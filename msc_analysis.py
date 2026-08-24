
import argparse, glob, json, os, re, sys
from collections import defaultdict
from math import comb

import numpy as np
import pandas as pd

DEFAULT_ROOTS = ["results", "results_binary", "results_cnad", "results_dro", "runs"]
DEFAULT_OUT = os.path.join("analysis_out", "final")
AGE_EDGES = [(0, 70, "<70"), (70, 80, "70-80"), (80, 200, "80+")]
N_BOOT = 2000
SEED = 1337


# OUTPUT HELPERS

def say(m):
    print(m); sys.stdout.flush()

def warn(m):
    print("WARNING: {}".format(m), file=sys.stderr)

def die(m):
    print("ERROR: {}".format(m), file=sys.stderr); sys.exit(1)

def banner(t):
    say("\n" + "-" * 10); say(t); say("-" * 10)

def write(df, outdir, name, force=False):
    os.makedirs(outdir, exist_ok=True)
    p = os.path.join(outdir, name)
    if os.path.exists(p) and not force:
        p = p.replace(".csv", "_new.csv")
        say("  (existing file kept; pass --force to overwrite in place)")
    df.to_csv(p, index=False)
    say("  written: {}  ({} rows)".format(p, len(df)))
    return p

def fmt(v):
    return "{:.3f}".format(v) if isinstance(v, float) else str(v)

def show(df, cols=None, n=200):
    d = df[cols] if cols else df
    say(d.head(n).to_string(index=False, float_format=lambda v: "{:.3f}".format(v)))


# =========================================================================
# DISCOVERY
# =========================================================================

def find_predictions(roots):
    hits = []
    for root in roots:
        if not os.path.isdir(root):
            continue
        for dp, _dn, fns in os.walk(root):
            for fn in fns:
                if fn.lower() == "predictions_test.csv":
                    hits.append(os.path.join(dp, fn))
    return sorted(hits)


def cohort_of(p):
    s = str(p).replace("\\", "/").lower()
    if "adni" in s:
        return "adni"
    if "oasis" in s:
        return "oasis"
    return None


def _run_dir_of(pred_path):
    """
    The run directory is where best.pt / train_summary.json live. Predictions
    usually sit one level below it in eval_within/ or eval_cross/.
    """
    d = os.path.dirname(pred_path)
    for cand in (d, os.path.dirname(d)):
        if any(os.path.isfile(os.path.join(cand, f))
               for f in ("train_summary.json", "best.pt")):
            return cand
    return os.path.dirname(d) or d


def parse_cell(path):
    """
    Identify a run.

    The training source comes from train_summary.json in the RUN directory
    (args.splits_dir there is what the model was trained on). The evaluation
    target comes from the evaluation summary beside the predictions, whose
    splits_dir is the cohort the model was APPLIED to. Reading the evaluation
    summary's splits_dir as the training source swaps every cross-cohort row,
    so the two are kept strictly separate here.

    Grouping keys are never left as None: pandas .groupby() silently drops NaN
    keys, which would delete a whole cell from every summary without an error.
    """
    info = dict(path=path, arm=None, backbone=None, train_cohort=None,
                eval_cohort=None, fold=None, regime=None, dro=False)
    eval_dir = os.path.dirname(path)
    run_dir = _run_dir_of(path)

    # ---- training side: train_summary.json in the run directory -------------
    ts = os.path.join(run_dir, "train_summary.json")
    if os.path.isfile(ts):
        try:
            js = json.load(open(ts))
            cfg = js.get("cfg") or {}
            args = js.get("args") or {}
            info["backbone"] = cfg.get("backbone") or args.get("backbone")
            sd = args.get("splits_dir") or js.get("splits_dir")
            if sd:
                info["train_cohort"] = cohort_of(sd)
            if cfg.get("out_dim"):
                info["arm"] = "2class" if int(cfg["out_dim"]) == 2 else "3class"
            if args.get("group_dro"):
                info["dro"] = True
        except Exception:
            pass

    # ---- evaluation side: the summary beside the predictions ----------------
    for cand in ("evaluation_summary.json", "eval_summary.json", "summary.json"):
        pjs = os.path.join(eval_dir, cand)
        if os.path.isfile(pjs):
            try:
                js = json.load(open(pjs))
                for k in ("eval_splits_dir", "target_splits_dir", "eval_dir", "splits_dir"):
                    if js.get(k):
                        info["eval_cohort"] = cohort_of(js[k]); break
                if info["backbone"] is None:
                    info["backbone"] = js.get("arch") or js.get("backbone")
                if info["arm"] is None and js.get("num_classes"):
                    info["arm"] = "2class" if int(js["num_classes"]) == 2 else "3class"
                # only trust a training dir if it is explicitly labelled as such
                if info["train_cohort"] is None:
                    for k in ("train_splits_dir", "train_dir"):
                        if js.get(k):
                            info["train_cohort"] = cohort_of(js[k]); break
            except Exception:
                pass
            break

    low = path.replace("\\", "/").lower()
    run_low = run_dir.replace("\\", "/").lower()
    eval_low = os.path.basename(eval_dir).lower()

    if info["fold"] is None:
        m = re.search(r"fold[_-]?(\d+)", run_low) or re.search(r"fold[_-]?(\d+)", low)
        if m:
            info["fold"] = int(m.group(1))
    if info["backbone"] is None:
        for k, v in (("medicalnet", "medicalnet"), ("mobilenet", "mobilenet3d_v2"),
                     ("r3d_18", "r3d_18"), ("r3d", "r3d_18")):
            if k in run_low:
                info["backbone"] = v; break
    if info["train_cohort"] is None:
        info["train_cohort"] = cohort_of(run_low)
    if info["arm"] is None:
        info["arm"] = "2class" if any(k in run_low for k in ("cnad", "binary", "bin_")) else "3class"
    if "dro" in run_low:
        info["dro"] = True

    # ---- resolve the evaluation target -------------------------------------
    other = {"adni": "oasis", "oasis": "adni"}
    if info["eval_cohort"] is None:
        if "cross" in eval_low or "cross" in low:
            info["eval_cohort"] = other.get(info["train_cohort"])
        elif "within" in eval_low or info["train_cohort"]:
            info["eval_cohort"] = info["train_cohort"]

    # a directory explicitly named cross must not come back as within
    if ("cross" in eval_low and info["eval_cohort"] == info["train_cohort"]
            and info["train_cohort"] in other):
        info["eval_cohort"] = other[info["train_cohort"]]

    # If the training cohort could not be resolved but the evaluation cohort
    # could, the eval directory name settles it: an eval_within directory was
    # evaluated on the training cohort, an eval_cross directory on the other.
    if info["train_cohort"] is None and info["eval_cohort"] in other:
        if "within" in eval_low:
            info["train_cohort"] = info["eval_cohort"]
        elif "cross" in eval_low:
            info["train_cohort"] = other[info["eval_cohort"]]
    # and symmetrically
    if info["eval_cohort"] is None and info["train_cohort"] in other:
        info["eval_cohort"] = (other[info["train_cohort"]] if "cross" in eval_low
                               else info["train_cohort"])

    info["regime"] = ("within" if (info["eval_cohort"] == info["train_cohort"]
                                   and info["eval_cohort"]) else "cross")
    for k in ("arm", "backbone", "train_cohort", "eval_cohort", "regime"):
        if info[k] is None:
            info[k] = "unknown"
    if info["fold"] is None:
        info["fold"] = -1
    return info


def load_predictions(path):
    """Normalise column names. Fails loudly rather than guessing."""
    df = pd.read_csv(path)
    cols = {c.lower().strip(): c for c in df.columns}

    def pick(*names):
        for n in names:
            if n in cols:
                return cols[n]
        return None

    c_sub = pick("subject_id", "subject", "sub_id", "id", "patient_id", "oasis_id", "rid")
    c_age = pick("age", "age_years", "scan_age", "age_at_scan")
    c_y = pick("label", "true_label", "y_true", "target", "y", "gt")
    c_p = pick("pred", "prediction", "y_pred", "pred_label", "predicted")
    miss = [n for n, c in (("age", c_age), ("label", c_y), ("pred", c_p)) if c is None]
    if miss:
        raise ValueError("{}: missing {}. Columns: {}".format(path, miss, list(df.columns)))

    probs = [cols[c] for c in sorted(cols) if re.fullmatch(r"prob_?\d+", c)]
    out = pd.DataFrame({
        "subject": df[c_sub].astype(str) if c_sub else np.arange(len(df)).astype(str),
        "age": pd.to_numeric(df[c_age], errors="coerce"),
        "y": pd.to_numeric(df[c_y], errors="coerce"),
        "pred": pd.to_numeric(df[c_p], errors="coerce"),
    })
    for i, pc in enumerate(probs):
        out["prob_{}".format(i)] = pd.to_numeric(df[pc], errors="coerce")
    for extra in ("dataset", "source", "collection", "sex", "session_id"):
        c = cols.get(extra)
        if c is not None:
            out[extra] = df[c]
    out = out.dropna(subset=["age", "y", "pred"])
    out["y"] = out["y"].astype(int); out["pred"] = out["pred"].astype(int)
    return out


def load_all(roots, quiet=False):
    """Return a list of (info, dataframe). Reports coverage."""
    files = find_predictions(roots)
    if not files:
        die("No predictions_test.csv under {}.\nRun from the project root, or pass "
            "--roots.".format(roots))
    recs = []
    for f in files:
        try:
            recs.append((parse_cell(f), load_predictions(f)))
        except ValueError as e:
            warn(str(e))
    if not quiet:
        cells = defaultdict(set)
        for i, _ in recs:
            cells[(i["arm"], i["backbone"], i["train_cohort"], i["eval_cohort"])].add(i["fold"])
        say("Loaded {} prediction files across {} cells.".format(len(recs), len(cells)))
        short = {k: v for k, v in cells.items() if len(v) != 5}
        if short:
            say("\n{} cell(s) without exactly 5 folds:".format(len(short)))
            for k, v in sorted(short.items(), key=lambda kv: tuple(map(str, kv[0]))):
                say("  {:<7} {:<15} {:>5} -> {:<5}  folds: {}".format(*[str(x) for x in k],
                                                                     sorted(v)))
        say("")
    return recs


def probs_of(df):
    pc = [c for c in df.columns if c.startswith("prob_")]
    if not pc:
        return None
    P = df[pc].values.astype(float)
    if not np.allclose(P.sum(axis=1), 1.0, atol=0.05):
        e = np.exp(P - P.max(axis=1, keepdims=True))
        P = e / e.sum(axis=1, keepdims=True)
    return P


def assign_bins(ages, edges=None):
    edges = edges or AGE_EDGES
    lab = np.full(len(ages), "NA", dtype=object)
    a = np.asarray(ages, float)
    for lo, hi, nm in edges:
        lab[(a >= lo) & (a < hi)] = nm
    return lab


def pick_col(df, prefs):
    """
    Choose a column by PREFERENCE order, not by the order columns happen to
    appear in the file. Iterating over df.columns picks whichever candidate
    comes first in the CSV, which selected 'diagnosis' over 'label'.
    """
    low = {str(c).lower().strip(): c for c in df.columns}
    for want in prefs:
        if want in low:
            return low[want]
    return None


LABEL_PREFS = ["label", "y", "target", "class", "diagnosis", "dx"]
AGE_PREFS = ["age", "age_years", "scan_age", "age_at_scan"]
SUBJ_PREFS = ["subject_id", "subject", "sub_id", "id", "patient_id", "oasis_id", "rid"]
PATH_PREFS = ["scan_path", "path", "filepath", "file", "nii_path"]


def parse_edges(spec):
    if not spec:
        return AGE_EDGES
    cuts = [0] + [float(x) for x in spec.split(",")] + [200]
    out = []
    for i in range(len(cuts) - 1):
        lo, hi = cuts[i], cuts[i + 1]
        nm = ("<{:.0f}".format(hi) if i == 0 else "{:.0f}+".format(lo)
              if i == len(cuts) - 2 else "{:.0f}-{:.0f}".format(lo, hi))
        out.append((lo, hi, nm))
    return out


# =========================================================================
# METRICS
# =========================================================================

def rates(y, p, k):
    y = np.asarray(y); p = np.asarray(p)
    o = {n: np.full(k, np.nan) for n in
         ("sens", "spec", "fnr", "ppv", "npv", "f1", "support")}
    for c in range(k):
        tp = np.sum((y == c) & (p == c)); fn = np.sum((y == c) & (p != c))
        fp = np.sum((y != c) & (p == c)); tn = np.sum((y != c) & (p != c))
        o["support"][c] = tp + fn
        o["sens"][c] = tp / (tp + fn) if tp + fn else np.nan
        o["spec"][c] = tn / (tn + fp) if tn + fp else np.nan
        o["fnr"][c] = fn / (tp + fn) if tp + fn else np.nan
        o["ppv"][c] = tp / (tp + fp) if tp + fp else np.nan
        o["npv"][c] = tn / (tn + fn) if tn + fn else np.nan
        pr, rc = o["ppv"][c], o["sens"][c]
        o["f1"][c] = (2 * pr * rc / (pr + rc)) if (pr and rc and pr + rc) else 0.0
    return o


def macro_f1(y, p, k):
    return float(np.nanmean(rates(y, p, k)["f1"]))


def bal_acc(y, p, k):
    return float(np.nanmean(rates(y, p, k)["sens"]))


def accuracy(y, p):
    return float(np.mean(np.asarray(y) == np.asarray(p)))


def roc_auc(y, P, k):
    """One-versus-rest macro AUC. Reduces to binary AUC when k == 2."""
    y = np.asarray(y)
    if P is None:
        return np.nan
    if k == 2:
        return _auc_binary(y == 1, P[:, 1])
    vals = [_auc_binary(y == c, P[:, c]) for c in range(k) if (y == c).any() and (y != c).any()]
    return float(np.nanmean(vals)) if vals else np.nan


def _auc_binary(pos, score):
    pos = np.asarray(pos, bool); score = np.asarray(score, float)
    n1, n0 = pos.sum(), (~pos).sum()
    if n1 == 0 or n0 == 0:
        return np.nan
    r = pd.Series(score).rank().values
    return float((r[pos].sum() - n1 * (n1 + 1) / 2) / (n1 * n0))


def ece_binned(y, P, n_bins=10):
    """Size-weighted mean |accuracy - confidence| within confidence bins."""
    if P is None:
        return np.nan
    conf = P.max(axis=1); pred = P.argmax(axis=1)
    correct = (pred == np.asarray(y)).astype(float)
    edges = np.linspace(0, 1, n_bins + 1)
    tot = 0.0
    for i in range(n_bins):
        m = (conf > edges[i]) & (conf <= edges[i + 1])
        if m.sum() == 0:
            continue
        tot += m.sum() / len(conf) * abs(correct[m].mean() - conf[m].mean())
    return float(tot)


def degenerate(p):
    return len(np.unique(np.asarray(p))) == 1


def boot_ci(fn, n, rng, reps=N_BOOT, alpha=0.05):
    v = np.array([fn(rng.integers(0, n, n)) for _ in range(reps)])
    lo, hi = np.nanpercentile(v, [100 * alpha / 2, 100 * (1 - alpha / 2)])
    return float(np.nanmean(v)), float(lo), float(hi)


def sign_test(deltas):
    d = np.asarray(deltas); d = d[d != 0]; n = len(d)
    if n == 0:
        return np.nan
    pos = int((d > 0).sum()); k = max(pos, n - pos)
    return float(min(1.0, 2 * sum(comb(n, i) for i in range(k, n + 1)) / 2 ** n))


def group_cells(recs):
    """Group (info, df) records into cells keyed by arm/backbone/train/eval."""
    cells = defaultdict(list)
    for i, d in recs:
        cells[(i["arm"], i["backbone"], i["train_cohort"], i["eval_cohort"],
               i["regime"], i["dro"])].append((i, d))
    return cells


def cellname(k):
    arm, bk, tr, ev, rg, dro = k
    return "{} {} {}->{}{}".format(arm, bk, tr, ev, " [DRO]" if dro else "")


# =========================================================================
# COMMAND: index
# =========================================================================

def cmd_index(a):
    recs = load_all(a.roots)
    rows = []
    for i, d in recs:
        P = probs_of(d); k = int(max(d.y.max(), d.pred.max())) + 1
        rows.append(dict(arm=i["arm"], backbone=i["backbone"], train=i["train_cohort"],
                         eval=i["eval_cohort"], regime=i["regime"], dro=i["dro"],
                         fold=i["fold"], n=len(d), n_classes=k,
                         degenerate=degenerate(d.pred.values),
                         macro_f1=macro_f1(d.y.values, d.pred.values, k),
                         balanced_acc=bal_acc(d.y.values, d.pred.values, k),
                         accuracy=accuracy(d.y.values, d.pred.values),
                         roc_auc=roc_auc(d.y.values, P, k),
                         ece=ece_binned(d.y.values, P), path=i["path"]))
    df = pd.DataFrame(rows).sort_values(["arm", "backbone", "train", "eval", "fold"])
    write(df, a.out, "run_index.csv", a.force)
    say("\nEvery run, one row each. This is the master index the other commands use.")
    return df


# =========================================================================
# COMMAND: cv   (Tables S5, S6, S8, S9, S18)
# =========================================================================

def cmd_cv(a):
    recs = load_all(a.roots)
    rng = np.random.default_rng(SEED)
    rows = []
    for key, mem in sorted(group_cells(recs).items(), key=lambda kv: tuple(map(str, kv[0]))):
        arm, bk, tr, ev, rg, dro = key
        k = int(max(max(d.y.max(), d.pred.max()) for _, d in mem)) + 1
        per = []
        for _, d in mem:
            P = probs_of(d)
            per.append((macro_f1(d.y.values, d.pred.values, k),
                        bal_acc(d.y.values, d.pred.values, k),
                        roc_auc(d.y.values, P, k),
                        accuracy(d.y.values, d.pred.values),
                        ece_binned(d.y.values, P)))
        per = np.array(per, float)

        # Within cohort the folds partition the cohort, so pool subjects and
        # resample. Across cohorts every fold sees the same subjects, so pooling
        # would count each subject five times: resample once, recompute per fold,
        # average. This is the protocol in Section 3.5.
        if rg == "within":
            allp = pd.concat([d for _, d in mem], ignore_index=True)
            y, p, P = allp.y.values, allp.pred.values, probs_of(allp)
            n = len(allp)
            mf, mlo, mhi = boot_ci(lambda ix: macro_f1(y[ix], p[ix], k), n, rng, a.reps)
            af, alo, ahi = boot_ci(
                lambda ix: roc_auc(y[ix], P[ix] if P is not None else None, k), n, rng, a.reps)
        else:
            n = len(mem[0][1])
            ys = [d.y.values for _, d in mem]; ps = [d.pred.values for _, d in mem]
            Ps = [probs_of(d) for _, d in mem]
            mf, mlo, mhi = boot_ci(
                lambda ix: np.mean([macro_f1(y[ix], p[ix], k) for y, p in zip(ys, ps)]),
                n, rng, a.reps)
            af, alo, ahi = boot_ci(
                lambda ix: np.nanmean([roc_auc(y[ix], (P[ix] if P is not None else None), k)
                                       for y, P in zip(ys, Ps)]), n, rng, a.reps)
        rows.append(dict(
            arm=arm, backbone=bk, train=tr, eval=ev, regime=rg, dro=dro, n_folds=len(mem),
            macro_f1=per[:, 0].mean(), macro_f1_sd=per[:, 0].std(ddof=1),
            bal_acc=per[:, 1].mean(), bal_acc_sd=per[:, 1].std(ddof=1),
            roc_auc=per[:, 2].mean(), roc_auc_sd=per[:, 2].std(ddof=1),
            accuracy=per[:, 3].mean(), ece=per[:, 4].mean(),
            macro_f1_boot=mf, macro_f1_lo=mlo, macro_f1_hi=mhi,
            roc_auc_boot=af, roc_auc_lo=alo, roc_auc_hi=ahi,
            auc_above_chance=bool(alo > 0.5),
            n_degenerate_folds=sum(degenerate(d.pred.values) for _, d in mem)))
    df = pd.DataFrame(rows)
    write(df, a.out, "cv_performance.csv", a.force)
    say("\nAll cells, fold means with bootstrap intervals:")
    df["variant"] = np.where(df["dro"], "DRO", "base")
    show(df, ["arm", "backbone", "train", "eval", "variant", "macro_f1", "bal_acc",
              "roc_auc", "roc_auc_lo", "roc_auc_hi", "auc_above_chance"])
    nb = int((~df.auc_above_chance).sum())
    say("\n{} of {} cells have a ROC-AUC interval containing 0.500 and cannot be "
        "described as above chance.".format(nb, len(df)))

    # within-to-cross degradation by metric (Table S7)
    deg = []
    for (arm, bk, tr, dr), g in df.groupby(["arm", "backbone", "train", "dro"],
                                           dropna=False):
        w = g[g.regime == "within"]; c = g[g.regime == "cross"]
        if len(w) and len(c):
            dm = c.macro_f1.iloc[0] - w.macro_f1.iloc[0]
            da = c.roc_auc.iloc[0] - w.roc_auc.iloc[0]
            deg.append(dict(arm=arm, backbone=bk, train=tr,
                            variant=("DRO" if dr else "base"),
                            d_macro_f1=dm, d_roc_auc=da,
                            ratio=abs(dm) / abs(da) if abs(da) > 1e-9 else np.inf))
    if deg:
        dd = pd.DataFrame(deg)
        write(dd, a.out, "degradation_by_metric.csv", a.force)
        say("\nWithin-to-cross degradation. macro-F1 falls faster than ROC-AUC:")
        show(dd)
    return df


# =========================================================================
# COMMAND: paired   (Table S19)
# =========================================================================

def cmd_paired(a):
    recs = load_all(a.roots)
    rng = np.random.default_rng(SEED)
    by = defaultdict(dict)
    for i, d in recs:
        by[(i["arm"], i["train_cohort"], i["eval_cohort"], i["regime"])].setdefault(
            i["backbone"], []).append((i["fold"], d))
    rows = []
    for key, bks in sorted(by.items(), key=lambda kv: tuple(map(str, kv[0]))):
        names = sorted(bks)
        for x in range(len(names)):
            for yb in range(x + 1, len(names)):
                A = dict(bks[names[x]]); B = dict(bks[names[yb]])
                folds = sorted(set(A) & set(B))
                if not folds:
                    continue
                k = int(max(max(A[f].y.max(), A[f].pred.max()) for f in folds)) + 1
                n = min(len(A[f]) for f in folds)
                # cache arrays once; `probs_of(...) or ...` is a truth-test on an
                # ndarray and raises, so probabilities are resolved up front
                PA = {f: probs_of(A[f]) for f in folds}
                PB = {f: probs_of(B[f]) for f in folds}
                YA = {f: A[f].y.values for f in folds}
                QA = {f: A[f].pred.values for f in folds}
                YB = {f: B[f].y.values for f in folds}
                QB = {f: B[f].pred.values for f in folds}
                for metric in ("macro_f1", "roc_auc"):
                    diffs = np.empty(a.reps)
                    for r in range(a.reps):
                        ix = rng.integers(0, n, n)
                        if metric == "macro_f1":
                            va = np.mean([macro_f1(YA[f][ix], QA[f][ix], k) for f in folds])
                            vb = np.mean([macro_f1(YB[f][ix], QB[f][ix], k) for f in folds])
                        else:
                            va = np.nanmean([roc_auc(YA[f][ix],
                                                     None if PA[f] is None else PA[f][ix], k)
                                             for f in folds])
                            vb = np.nanmean([roc_auc(YB[f][ix],
                                                     None if PB[f] is None else PB[f][ix], k)
                                             for f in folds])
                        diffs[r] = va - vb
                    lo, hi = np.nanpercentile(diffs, [2.5, 97.5])
                    rows.append(dict(arm=key[0], train=key[1], eval=key[2], regime=key[3],
                                     metric=metric, a=names[x], b=names[yb],
                                     diff=float(np.nanmean(diffs)), lo=float(lo), hi=float(hi),
                                     excludes_zero=bool(lo > 0 or hi < 0)))
    df = pd.DataFrame(rows)
    write(df, a.out, "paired_backbone.csv", a.force)
    if len(df):
        say("\nPaired backbone comparisons on identical resamples.")
        say("{} of {} exclude zero. In the three-class arm: {} of {}.".format(
            int(df.excludes_zero.sum()), len(df),
            int(df[df.arm == "3class"].excludes_zero.sum()), int((df.arm == "3class").sum())))
        sig = df[df.excludes_zero]
        if len(sig):
            show(sig, ["arm", "train", "eval", "metric", "a", "b", "diff", "lo", "hi"])
    return df


# =========================================================================
# COMMAND: perbin   (proposal item: per-bin sensitivity / specificity / FNR)
# =========================================================================

def cmd_perbin(a):
    recs = load_all(a.roots)
    edges = parse_edges(a.age_edges)
    rng = np.random.default_rng(SEED)
    rows = []
    for i, d in recs:
        k = int(max(d.y.max(), d.pred.max())) + 1
        d = d.copy(); d["bin"] = assign_bins(d.age.values, edges)
        dg = degenerate(d.pred.values)
        for b in [e[2] for e in edges]:
            s = d[d.bin == b]
            if len(s) < a.min_bin:
                continue
            r = rates(s.y.values, s.pred.values, k)
            y, p = s.y.values, s.pred.values
            m, lo, hi = boot_ci(lambda ix: macro_f1(y[ix], p[ix], k), len(s), rng, a.reps)
            dis = k - 1
            rows.append(dict(
                arm=i["arm"], backbone=i["backbone"], train=i["train_cohort"],
                eval=i["eval_cohort"], regime=i["regime"], dro=i["dro"], fold=i["fold"],
                age_bin=b, n=len(s), degenerate_fold=dg,
                macro_f1=macro_f1(y, p, k), macro_f1_lo=lo, macro_f1_hi=hi,
                balanced_acc=bal_acc(y, p, k),
                sens_disease=r["sens"][dis], spec_disease=r["spec"][dis],
                fnr_disease=r["fnr"][dis], ppv_disease=r["ppv"][dis],
                npv_disease=r["npv"][dis], support_disease=r["support"][dis]))
    df = pd.DataFrame(rows)
    write(df, a.out, "per_bin_metrics.csv", a.force)
    if len(df):
        agg = df.groupby(["arm", "backbone", "train", "eval", "dro", "age_bin"],
                         dropna=False).agg(
            n_folds=("fold", "nunique"), any_degenerate=("degenerate_fold", "any"),
            sens_disease=("sens_disease", "mean"), sens_sd=("sens_disease", "std"),
            spec_disease=("spec_disease", "mean"), fnr_disease=("fnr_disease", "mean"),
            ppv_disease=("ppv_disease", "mean"), npv_disease=("npv_disease", "mean"),
            macro_f1=("macro_f1", "mean")).reset_index()
        write(agg, a.out, "per_bin_metrics_foldmean.csv", a.force)
        say("\nPer-bin disease-class rates (fold means).")
        say("any_degenerate=True: a fold predicted one class for everyone, so the row")
        say("is not an age effect. n_folds<5: the bin fell under --min-bin ({}).\n".format(a.min_bin))
        agg["variant"] = np.where(agg["dro"], "DRO", "base")
        show(agg, ["arm", "backbone", "train", "eval", "variant", "age_bin", "n_folds",
                   "any_degenerate", "sens_disease", "spec_disease", "fnr_disease"])
    return df


# =========================================================================
# COMMAND: bins   (proposal item: alternative cut-points / quantile bins)
# =========================================================================

def cmd_bins(a):
    recs = load_all(a.roots)
    schemes = {"fixed_70_80": AGE_EDGES,
               "fixed_65_75_85": [(0, 65, "<65"), (65, 75, "65-75"),
                                  (75, 85, "75-85"), (85, 200, "85+")],
               "fixed_75": [(0, 75, "<75"), (75, 200, "75+")],
               "tertiles": None, "quartiles": None, "median": None}
    rows = []
    for key, mem in sorted(group_cells(recs).items(), key=lambda kv: tuple(map(str, kv[0]))):
        allp = pd.concat([d.assign(fold=i["fold"]) for i, d in mem], ignore_index=True)
        k = int(max(allp.y.max(), allp.pred.max())) + 1
        ages = allp.age.values
        for sn, edges in schemes.items():
            if edges is None:
                q = {"tertiles": [1 / 3, 2 / 3], "quartiles": [.25, .5, .75],
                     "median": [.5]}[sn]
                cuts = [0] + [float(x) for x in np.quantile(ages, q)] + [200]
                edges = [(cuts[j], cuts[j + 1], "q{}".format(j + 1))
                         for j in range(len(cuts) - 1)]
            gaps, ndeg = [], 0
            for _f, sub in allp.groupby("fold"):
                if degenerate(sub.pred.values):
                    ndeg += 1; continue
                sub = sub.copy(); sub["bin"] = assign_bins(sub.age.values, edges)
                v = [macro_f1(s.y.values, s.pred.values, k)
                     for b in [e[2] for e in edges]
                     for s in [sub[sub.bin == b]] if len(s) >= a.min_bin]
                if len(v) >= 2:
                    gaps.append(max(v) - min(v))
            rows.append(dict(arm=key[0], backbone=key[1], train=key[2], eval=key[3],
                             variant=("DRO" if key[5] else "base"),
                             scheme=sn, n_bins=len(edges),
                             gap_mean=float(np.mean(gaps)) if gaps else np.nan,
                             gap_sd=float(np.std(gaps, ddof=1)) if len(gaps) > 1 else np.nan,
                             n_folds_used=len(gaps), n_degenerate=ndeg,
                             usable=len(gaps) >= 2))
    df = pd.DataFrame(rows)
    write(df, a.out, "age_bin_sensitivity.csv", a.force)
    good = df[df.usable]
    say("\nWorst-minus-best age gap under each binning. Stability across rows means")
    say("the gap is not an artefact of the 70/80 cut-points.\n")
    if len(good):
        piv = good.pivot_table(index=["arm", "backbone", "train", "eval", "variant"],
                               columns="scheme", values="gap_mean")
        say(piv.to_string(float_format=lambda v: "{:.3f}".format(v)))
        sp = good.groupby(["arm", "backbone", "train", "eval", "variant"])["gap_mean"].agg(["min", "max"])
        sp["range"] = sp["max"] - sp["min"]
        say("\nRange across schemes (small = robust to bin choice):")
        say(sp[["range"]].to_string(float_format=lambda v: "{:.3f}".format(v)))
    else:
        say("  No cell has two or more non-degenerate folds; see usable=False rows.")
    return df


# =========================================================================
# COMMAND: confage   (proposal item: confidence-age correlation)
# =========================================================================

def cmd_confage(a):
    from scipy import stats
    recs = load_all(a.roots)
    rng = np.random.default_rng(SEED)
    rows = []
    for i, d in recs:
        P = probs_of(d)
        if P is None:
            continue
        conf = P.max(axis=1); dis = P[:, -1]; age = d.age.values.astype(float)
        corr = (d.y.values == d.pred.values).astype(float)

        def c(x, y):
            if np.std(x) == 0 or np.std(y) == 0:
                return np.nan, np.nan
            return stats.pearsonr(x, y)

        r1, p1 = c(age, conf); r2, p2 = c(age, dis); r3, p3 = c(age, corr)
        pp = np.nan
        if not np.isnan(r2):
            obs = abs(r2)
            pp = (sum(abs(np.corrcoef(rng.permutation(age), dis)[0, 1]) >= obs
                      for _ in range(a.perm)) + 1) / (a.perm + 1)
        rows.append(dict(arm=i["arm"], backbone=i["backbone"], train=i["train_cohort"],
                         eval=i["eval_cohort"], regime=i["regime"],
                         variant=("DRO" if i["dro"] else "base"),
                         fold=i["fold"], n=len(d),
                         degenerate_fold=degenerate(d.pred.values),
                         r_age_conf=r1, p_age_conf=p1,
                         r_age_disease_prob=r2, p_age_disease_prob=p2, perm_p=pp,
                         r_age_correct=r3, p_age_correct=p3, mean_conf=float(conf.mean())))
    df = pd.DataFrame(rows)
    write(df, a.out, "confidence_age.csv", a.force)
    if len(df):
        agg = df[~df.degenerate_fold].groupby(
            ["arm", "backbone", "train", "eval", "variant"], dropna=False).agg(
            n_folds=("fold", "nunique"), r_age_conf=("r_age_conf", "mean"),
            r_age_disease=("r_age_disease_prob", "mean"),
            r_age_disease_sd=("r_age_disease_prob", "std")).reset_index()
        write(agg, a.out, "confidence_age_foldmean.csv", a.force)
        say("\nAge versus confidence and versus the disease-class posterior.")
        say("Positive r_age_disease is the quantitative form of age-tracking.\n")
        show(agg)
    return df


# =========================================================================
# COMMAND: thresh   (Tables S14, S15, S16 + temperature scaling)
# =========================================================================

def _prevalence_match(P, target_prior):
    """Label so predicted class counts match the target proportions."""
    n, k = P.shape
    want = np.round(np.asarray(target_prior) * n).astype(int)
    want[-1] = n - want[:-1].sum()
    out = np.full(n, -1)
    order = np.argsort(-P.max(axis=1))
    left = want.copy()
    for idx in order:
        for c in np.argsort(-P[idx]):
            if left[c] > 0:
                out[idx] = c; left[c] -= 1; break
        if out[idx] == -1:
            out[idx] = int(np.argmax(P[idx]))
    return out


def _prior_correct(P, train_prior, target_prior):
    eps = 1e-12
    adj = P * (np.asarray(target_prior) + eps) / (np.asarray(train_prior) + eps)
    return adj.argmax(axis=1)


def _fit_temperature(P, y, grid=None):
    grid = grid if grid is not None else np.concatenate([np.linspace(0.05, 5, 100),
                                                         np.linspace(5.5, 50, 90)])
    logits = np.log(np.clip(P, 1e-12, 1))
    best, bt = np.inf, 1.0
    for t in grid:
        q = np.exp(logits / t); q /= q.sum(axis=1, keepdims=True)
        nll = -np.mean(np.log(np.clip(q[np.arange(len(y)), y], 1e-12, 1)))
        if nll < best:
            best, bt = nll, t
    return float(bt)


def cmd_thresh(a):
    recs = load_all(a.roots)
    rows = []
    for key, mem in sorted(group_cells(recs).items(), key=lambda kv: tuple(map(str, kv[0]))):
        k = int(max(max(d.y.max(), d.pred.max()) for _, d in mem)) + 1
        base, prior, prev, orac, tempr = [], [], [], [], []
        for _i, d in mem:
            P = probs_of(d)
            if P is None:
                continue
            y = d.y.values
            tgt = np.bincount(y, minlength=k) / len(y)
            trn = np.full(k, 1.0 / k)          # training was class-balanced
            base.append((macro_f1(y, d.pred.values, k), bal_acc(y, d.pred.values, k)))
            pc = _prior_correct(P, trn, tgt)
            prior.append((macro_f1(y, pc, k), bal_acc(y, pc, k), degenerate(pc)))
            pm = _prevalence_match(P, tgt)
            prev.append((macro_f1(y, pm, k), bal_acc(y, pm, k), degenerate(pm)))
            if k == 2:
                ths = np.unique(P[:, 1]); best = 0.0
                for t in ths:
                    best = max(best, bal_acc(y, (P[:, 1] >= t).astype(int), k))
                orac.append(best)
            t = _fit_temperature(P, y)
            lg = np.log(np.clip(P, 1e-12, 1)); q = np.exp(lg / t)
            q /= q.sum(axis=1, keepdims=True)
            tempr.append((t, ece_binned(y, q)))
        if not base:
            continue
        base = np.array(base); prior = np.array(prior, float); prev = np.array(prev, float)
        tempr = np.array(tempr, float)
        rows.append(dict(
            arm=key[0], backbone=key[1], train=key[2], eval=key[3], regime=key[4],
            variant=("DRO" if key[5] else "base"),
            base_macro_f1=base[:, 0].mean(), base_bal_acc=base[:, 1].mean(),
            prior_macro_f1=prior[:, 0].mean(), prior_bal_acc=prior[:, 1].mean(),
            prior_degenerate_folds=int(prior[:, 2].sum()),
            prev_macro_f1=prev[:, 0].mean(), prev_bal_acc=prev[:, 1].mean(),
            prev_degenerate_folds=int(prev[:, 2].sum()),
            oracle_bal_acc=float(np.mean(orac)) if orac else np.nan,
            mean_temperature=tempr[:, 0].mean(), ece_after_temp=tempr[:, 1].mean(),
            d_prior=prior[:, 0].mean() - base[:, 0].mean(),
            d_prev=prev[:, 0].mean() - base[:, 0].mean()))
    df = pd.DataFrame(rows)
    write(df, a.out, "threshold_interventions.csv", a.force)
    if len(df):
        say("\nThreshold interventions, macro-F1:")
        show(df, ["arm", "backbone", "train", "eval", "variant", "base_macro_f1",
                  "prior_macro_f1", "prev_macro_f1", "oracle_bal_acc", "d_prior", "d_prev"])
        say("\nPrevalence matching improves {} of {} cells; prior correction improves {}."
            .format(int((df.d_prev > 0).sum()), len(df), int((df.d_prior > 0).sum())))
        say("Prior correction produces {} degenerate folds; prevalence matching {}."
            .format(int(df.prior_degenerate_folds.sum()), int(df.prev_degenerate_folds.sum())))
        say("Mean fitted temperature {:.1f}. Above ~20 means confidence is being flattened "
            "toward uniform rather than corrected.".format(df.mean_temperature.mean()))
    return df


# =========================================================================
# COMMAND: agegap   (Table S21, the primary success criterion)
# =========================================================================

def cmd_agegap(a):
    recs = load_all(a.roots)
    edges = parse_edges(a.age_edges)
    rng = np.random.default_rng(SEED)
    names = [e[2] for e in edges]

    def gap(y, p, ages, k):
        v = []
        for b in names:
            m = assign_bins(ages, edges) == b
            if m.sum() >= a.min_bin:
                v.append(macro_f1(y[m], p[m], k))
        return (max(v) - min(v)) if len(v) >= 2 else np.nan

    rows = []
    for key, mem in sorted(group_cells(recs).items(), key=lambda kv: tuple(map(str, kv[0]))):
        if any(degenerate(d.pred.values) for _, d in mem):
            continue                                # excluded by protocol
        k = int(max(max(d.y.max(), d.pred.max()) for _, d in mem)) + 1
        for iv in ("prevalence", "prior"):
            diffs, b0, b1 = np.empty(a.reps), [], []
            n = min(len(d) for _, d in mem)
            for r in range(a.reps):
                ix = rng.integers(0, n, n)
                gb, ga = [], []
                for _i, d in mem:
                    P = probs_of(d)
                    if P is None:
                        continue
                    y = d.y.values[ix]; ages = d.age.values[ix]; Pi = P[ix]
                    tgt = np.bincount(y, minlength=k) / len(y)
                    alt = (_prevalence_match(Pi, tgt) if iv == "prevalence"
                           else _prior_correct(Pi, np.full(k, 1.0 / k), tgt))
                    gb.append(gap(y, d.pred.values[ix], ages, k))
                    ga.append(gap(y, alt, ages, k))
                # a resample can leave every bin under --min-bin; that draw
                # contributes nothing rather than emitting a NaN warning
                if np.all(np.isnan(ga)) or np.all(np.isnan(gb)):
                    diffs[r] = np.nan
                else:
                    diffs[r] = np.nanmean(ga) - np.nanmean(gb)
                if r == 0:
                    b0, b1 = np.nanmean(gb), np.nanmean(ga)
            lo, hi = np.nanpercentile(diffs, [2.5, 97.5])
            # Two-sided bootstrap p: twice the smaller tail mass about zero.
            d_ok = diffs[~np.isnan(diffs)]
            if len(d_ok):
                frac = float(np.mean(d_ok >= 0))
                p_boot = float(min(1.0, 2 * min(frac, 1 - frac)))
                p_boot = max(p_boot, 1.0 / (len(d_ok) + 1))
            else:
                p_boot = np.nan
            rows.append(dict(arm=key[0], backbone=key[1], train=key[2], eval=key[3],
                             regime=key[4], variant=("DRO" if key[5] else "base"),
                             intervention=iv, p_boot=p_boot,
                             gap_baseline=float(b0), gap_after=float(b1),
                             difference=float(np.nanmean(diffs)), lo=float(lo), hi=float(hi),
                             significant=bool(hi < 0 or lo > 0),
                             direction=("improves" if np.nanmean(diffs) < 0 else "worsens")))
    df = pd.DataFrame(rows)

    # ---- multiplicity ------------------------------------------------------
    # The success criterion is "significant in at least one cell". Testing ~40
    # cell-by-intervention combinations at a nominal 5 per cent and declaring
    # success on any hit inflates the family-wise error rate badly: under
    # independent nulls, 1 - 0.95^40 is about 0.87. Holm-Bonferroni is applied
    # across the whole family so the criterion can be stated as significant
    # after correction, not merely nominally significant.
    if len(df) and df.p_boot.notna().any():
        m = int(df.p_boot.notna().sum())
        order = df.p_boot.rank(method="first").astype("Int64")
        df["holm_alpha"] = 0.05 / (m - order + 1)
        d = df.dropna(subset=["p_boot"]).sort_values("p_boot")
        passed, thresh = [], True
        for i, (_ix, r) in enumerate(d.iterrows()):
            thresh = thresh and (r.p_boot <= 0.05 / (m - i))
            passed.append(thresh)
        df["holm_significant"] = False
        df.loc[d.index, "holm_significant"] = passed
        df["n_tests_in_family"] = m
    write(df, a.out, "age_gap_paired.csv", a.force)
    if len(df):
        say("\nPaired bootstrap on the worst-minus-best age-bin gap.")
        say("Negative difference = the gap narrowed. Cells with degenerate folds excluded.\n")
        cols = ["arm", "backbone", "train", "eval", "variant", "intervention",
                "gap_baseline", "gap_after", "difference", "lo", "hi", "significant"]
        if "holm_significant" in df.columns:
            cols += ["p_boot", "holm_significant"]
        show(df, cols)
        w = df[(df.significant) & (df.difference < 0)]
        say("\nNOMINAL: {} of {} cell-by-intervention tests show a reduction with a"
            .format(len(w), len(df)))
        say("95 per cent interval excluding zero, uncorrected.")
        if "holm_significant" in df.columns:
            hw = df[(df.holm_significant) & (df.difference < 0)]
            say("\nHOLM-BONFERRONI across the family of {} tests: {} survive."
                .format(int(df.n_tests_in_family.iloc[0]), len(hw)))
            if len(hw):
                say("  -> PRIMARY SUCCESS CRITERION MET after correction for multiplicity.")
                show(hw, ["arm", "backbone", "train", "eval", "intervention",
                          "difference", "lo", "hi", "p_boot"])
            else:
                say("  -> The criterion is met NOMINALLY but NOT after correction.")
                say("     Report it as exploratory evidence across cells, and say so.")
                say("     Testing many cells and declaring success on any hit inflates")
                say("     the family-wise error rate; with {} tests at alpha 0.05 the"
                    .format(int(df.n_tests_in_family.iloc[0])))
                say("     chance of at least one nominal hit under the null is high.")
        bad = df[(df.significant) & (df.difference > 0)]
        if len(bad):
            say("\n{} cell(s) significantly WORSEN and must be reported too."
                .format(len(bad)))
    return df


# =========================================================================
# COMMAND: degen   (Table S20)
# =========================================================================

def cmd_degen(a):
    recs = load_all(a.roots)
    rows = []
    for key, mem in sorted(group_cells(recs).items(), key=lambda kv: tuple(map(str, kv[0]))):
        k = int(max(max(d.y.max(), d.pred.max()) for _, d in mem)) + 1
        for iv in ("baseline", "prior", "prevalence"):
            nd, which = 0, []
            for i, d in mem:
                P = probs_of(d)
                if iv == "baseline":
                    p = d.pred.values
                elif P is None:
                    continue
                else:
                    tgt = np.bincount(d.y.values, minlength=k) / len(d)
                    p = (_prior_correct(P, np.full(k, 1.0 / k), tgt) if iv == "prior"
                         else _prevalence_match(P, tgt))
                if degenerate(p):
                    nd += 1; which.append(i["fold"])
            rows.append(dict(arm=key[0], backbone=key[1], train=key[2], eval=key[3],
                             regime=key[4], variant=("DRO" if key[5] else "base"),
                             intervention=iv, n_folds=len(mem),
                             n_degenerate=nd, all_folds_degenerate=(nd == len(mem)),
                             folds=",".join(map(str, sorted(which)))))
    df = pd.DataFrame(rows)
    write(df, a.out, "degeneracy_inventory.csv", a.force)
    bad = df[df.n_degenerate > 0]
    say("\n{} of {} cell-and-intervention combinations contain a degenerate fold."
        .format(len(bad), len(df)))
    if len(bad):
        say("By intervention: {}".format(bad.groupby("intervention").size().to_dict()))
        say("All five folds collapse in {} combination(s).".format(int(df.all_folds_degenerate.sum())))
        show(bad, ["arm", "backbone", "train", "eval", "variant", "intervention",
                   "n_degenerate", "folds"])
    return df


# =========================================================================
# COMMAND: domain   (Table S26)
# =========================================================================

def cmd_domain(a):
    """Stratify OASIS performance by source collection using a subject-id prefix."""
    recs = load_all(a.roots)
    rows = []
    for i, d in recs:
        if i["eval_cohort"] != "oasis":
            continue
        # Prefer an explicit source column if the prediction file carries one;
        # the split CSVs have `dataset` with values OAS1 / OAS2.
        src_col = next((c for c in d.columns
                        if str(c).lower() in ("dataset", "source", "collection")), None)
        if src_col is not None:
            is1 = d[src_col].astype(str).str.upper().str.startswith("OAS1")
        else:
            is1 = d.subject.astype(str).str.contains(a.oas1_pattern, case=False,
                                                     regex=True)
        if is1.sum() == 0 or (~is1).sum() == 0:
            continue
        k = int(max(d.y.max(), d.pred.max())) + 1
        P = probs_of(d)
        r = dict(arm=i["arm"], backbone=i["backbone"], train=i["train_cohort"],
                 regime=i["regime"], variant=("DRO" if i["dro"] else "base"),
                 fold=i["fold"],
                 n_oas1=int(is1.sum()), n_oas2=int((~is1).sum()),
                 auc_pooled=roc_auc(d.y.values, P, k))
        for nm, m in (("oas1", is1.values), ("oas2", (~is1).values)):
            r["auc_" + nm] = roc_auc(d.y.values[m], P[m] if P is not None else None, k)
            r["macro_f1_" + nm] = macro_f1(d.y.values[m], d.pred.values[m], k)
        r["domain_gap"] = r["auc_oas1"] - r["auc_oas2"]
        rows.append(r)
    df = pd.DataFrame(rows)
    if not len(df):
        say("\nNo OASIS evaluations matched --oas1-pattern ({}). Set it to whatever "
            "distinguishes OASIS-1 subject ids in your files.".format(a.oas1_pattern))
        return df
    write(df, a.out, "domain_stratified.csv", a.force)
    agg = df.groupby(["arm", "backbone", "train", "variant"], dropna=False).agg(
        auc_pooled=("auc_pooled", "mean"), auc_oas1=("auc_oas1", "mean"),
        auc_oas2=("auc_oas2", "mean"), domain_gap=("domain_gap", "mean")).reset_index()
    write(agg, a.out, "domain_stratified_foldmean.csv", a.force)
    say("\nOASIS-1 versus OASIS-2. The pooled AUC sits BELOW both subsets because")
    say("cross-subset pairs are ordered by pipeline rather than by disease.\n")
    show(agg)
    say("\nOASIS-1 higher in {} of {} cells.".format(int((agg.domain_gap > 0).sum()), len(agg)))
    return df


# =========================================================================
# COMMAND: calib   (Table S22 + Figure 4)
# =========================================================================

def cmd_calib(a):
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    recs = load_all(a.roots)
    edges = parse_edges(a.age_edges); names = [e[2] for e in edges]
    rows = []
    for i, d in recs:
        P = probs_of(d)
        if P is None:
            continue
        b = assign_bins(d.age.values, edges)
        r = dict(arm=i["arm"], backbone=i["backbone"], train=i["train_cohort"],
                 eval=i["eval_cohort"], regime=i["regime"],
                 variant=("DRO" if i["dro"] else "base"), fold=i["fold"],
                 degenerate_fold=degenerate(d.pred.values))
        for nm in names:
            m = b == nm
            r["ece_" + nm] = ece_binned(d.y.values[m], P[m]) if m.sum() >= a.min_bin else np.nan
        vals = [r["ece_" + n] for n in names if not np.isnan(r.get("ece_" + n, np.nan))]
        r["ece_gap"] = (max(vals) - min(vals)) if len(vals) >= 2 else np.nan
        r["worst_bin"] = (names[int(np.nanargmax([r.get("ece_" + n, np.nan) for n in names]))]
                          if len(vals) >= 2 else "")
        rows.append(r)
    df = pd.DataFrame(rows)
    write(df, a.out, "calibration_by_age.csv", a.force)
    if len(df):
        cell = df.groupby(["arm", "backbone", "train", "eval", "variant"],
                          dropna=False).agg(
            **{("ece_" + n): ("ece_" + n, "mean") for n in names}).reset_index()
        for c in cell.index:
            v = [cell.loc[c, "ece_" + n] for n in names]
            cell.loc[c, "worst_bin"] = names[int(np.nanargmax(v))] if not np.all(np.isnan(v)) else ""
        write(cell, a.out, "calibration_by_age_cell.csv", a.force)
        say("\nMean binned ECE by age bin:")
        for n in names:
            say("  {:>7}  {:.3f}   worst in {} of {} cells".format(
                n, cell["ece_" + n].mean(), int((cell.worst_bin == n).sum()), len(cell)))
        k = int((cell.worst_bin == names[0]).sum()); N = len(cell); m = len(names)
        p = sum(comb(N, j) * (1 / m) ** j * (1 - 1 / m) ** (N - j) for j in range(k, N + 1))
        say("\nYoungest bin worst in {} of {} against {:.1f} expected by chance, "
            "binomial p = {:.2g}".format(k, N, N / m, p))
        say("TRAP: the lowest ECE in the study may belong to a degenerate cell. A constant")
        say("predictor at the base rate is perfectly calibrated and useless. Check the flag.")

        # Figure 4: reliability by age bin
        fig, axes = plt.subplots(1, len(names), figsize=(4 * len(names), 3.6), sharey=True)
        for ax, nm in zip(np.atleast_1d(axes), names):
            xs, ys = [], []
            for i, d in recs:
                P = probs_of(d)
                if P is None or i["regime"] != "within":
                    continue
                m = assign_bins(d.age.values, edges) == nm
                if m.sum() < a.min_bin:
                    continue
                conf = P[m].max(axis=1); corr = (P[m].argmax(1) == d.y.values[m]).astype(float)
                bs = np.linspace(0, 1, 11)
                for j in range(10):
                    s = (conf > bs[j]) & (conf <= bs[j + 1])
                    if s.sum():
                        xs.append(conf[s].mean()); ys.append(corr[s].mean())
            ax.plot([0, 1], [0, 1], "k--", lw=1)
            ax.scatter(xs, ys, s=14, alpha=0.6)
            ax.set_title("age {}".format(nm), fontsize=10)
            ax.set_xlabel("confidence"); ax.set_xlim(0, 1); ax.set_ylim(0, 1)
        np.atleast_1d(axes)[0].set_ylabel("observed accuracy")
        fig.suptitle("Figure 4. Calibration by age bin (within-cohort)", fontsize=11)
        fig.tight_layout()
        os.makedirs(a.out, exist_ok=True)
        fig.savefig(os.path.join(a.out, "figure4_calibration.png"), dpi=160)
        plt.close(fig)
        say("  written: {}".format(os.path.join(a.out, "figure4_calibration.png")))
    return df


# =========================================================================
# COMMAND: figures   (Figures 2 and 3)
# =========================================================================

def cmd_figures(a):
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    recs = load_all(a.roots)
    os.makedirs(a.out, exist_ok=True)

    # Figure 2: age distribution by class and cohort
    fig, axes = plt.subplots(1, 2, figsize=(11, 3.8), sharey=True)
    for ax, coh in zip(axes, ["oasis", "adni"]):
        pool = [d for i, d in recs if i["train_cohort"] == coh and i["regime"] == "within"]
        if not pool:
            ax.set_visible(False); continue
        allp = pd.concat(pool, ignore_index=True).drop_duplicates("subject")
        for c in sorted(allp.y.unique()):
            ax.hist(allp[allp.y == c].age, bins=18, alpha=0.55,
                    label="class {} (n={})".format(c, int((allp.y == c).sum())))
        ax.set_title(coh.upper()); ax.set_xlabel("age"); ax.legend(fontsize=8)
    axes[0].set_ylabel("subjects")
    fig.suptitle("Figure 2. Age distribution by diagnostic class", fontsize=11)
    fig.tight_layout(); fig.savefig(os.path.join(a.out, "figure2_age_distribution.png"), dpi=160)
    plt.close(fig); say("  written: figure2_age_distribution.png")

    # Figure 3: predicted class counts per fold vs truth
    tgt = None
    for key, mem in group_cells(recs).items():
        if key[0] == "2class" and key[1] == "r3d_18" and key[2] == "adni" and key[4] == "cross":
            tgt = (key, mem); break
    if tgt:
        key, mem = tgt
        mem = sorted(mem, key=lambda m: m[0]["fold"])
        k = int(max(max(d.y.max(), d.pred.max()) for _, d in mem)) + 1
        fig, ax = plt.subplots(figsize=(7.5, 3.8))
        w = 0.8 / (k + 1)
        x = np.arange(len(mem))
        truth = np.bincount(mem[0][1].y.values, minlength=k)
        for c in range(k):
            ax.bar(x + c * w, [np.sum(d.pred.values == c) for _, d in mem], w,
                   label="predicted class {}".format(c))
        for c in range(k):
            ax.axhline(truth[c], ls="--", lw=1, color="grey")
        ax.set_xticks(x + w * (k - 1) / 2)
        ax.set_xticklabels(["fold {}".format(i["fold"]) for i, _ in mem])
        ax.set_ylabel("subjects"); ax.legend(fontsize=8)
        ax.set_title("Figure 3. Predicted class counts per fold against the fixed true "
                     "distribution\n({})".format(cellname(key)), fontsize=10)
        fig.tight_layout(); fig.savefig(os.path.join(a.out, "figure3_collapse.png"), dpi=160)
        plt.close(fig); say("  written: figure3_collapse.png")
    else:
        say("  Figure 3 skipped: could not find the ADNI->OASIS two-class R3D-18 cell.")


# =========================================================================
# TORCH-BACKED COMMANDS. torch is imported inside each function, never at
# module level, so a broken torch install cannot stop the CPU analyses.
# =========================================================================

def _torch():
    try:
        import torch
        return torch
    except ImportError:
        die("This command needs torch:  pip install torch")


def find_checkpoints(roots):
    out = []
    for r in roots:
        if os.path.isdir(r):
            out += glob.glob(os.path.join(r, "**", "*.pt"), recursive=True)
    return sorted(out)


# ---------------------------------------------------------------- config
def cmd_config(a):
    """Read the stored training arguments out of every checkpoint (Table S17)."""
    torch = _torch()
    cks = find_checkpoints(a.roots)
    if not cks:
        die("No .pt checkpoints under {}".format(a.roots))
    rows = []
    for c in cks:
        try:
            ck = torch.load(c, map_location="cpu", weights_only=False)
        except Exception as e:
            warn("{}: {}".format(c, e)); continue
        if not isinstance(ck, dict):
            continue
        args = ck.get("args", {}) or {}
        cfg = ck.get("cfg", {}) or {}
        rows.append(dict(path=c, epoch=ck.get("epoch"), monitor=ck.get("monitor"),
                         best_score_VALIDATION=ck.get("best_score"),
                         num_classes=ck.get("num_classes") or cfg.get("out_dim"),
                         backbone=cfg.get("backbone") or args.get("backbone"),
                         **{k: args.get(k) for k in
                            ("lr", "batch_size", "accum_steps", "patch_size", "dropout",
                             "weight_decay", "epochs", "freeze_epochs", "early_patience",
                             "label_smoothing", "focal_gamma", "ema", "balanced_sampler",
                             "no_class_weights", "group_dro", "dro_eta", "seed",
                             "splits_dir")}))
    df = pd.DataFrame(rows)
    write(df, a.out, "checkpoint_config.csv", a.force)
    say("\nSettings across {} checkpoints. Constant settings are claims you can make; "
        "varying ones are confounds you must declare.\n".format(len(df)))
    for c in ("ema", "balanced_sampler", "no_class_weights", "lr", "batch_size",
              "patch_size", "dropout", "monitor", "group_dro"):
        if c in df.columns:
            vc = df[c].value_counts(dropna=False).to_dict()
            flag = "" if len(vc) == 1 else "   <-- VARIES"
            say("  {:20s} {}{}".format(c, vc, flag))
    say("\nWarning: best_score_VALIDATION is a best-epoch validation score. It must "
        "never be reported as a result.")
    return df


# embed
def _find_train_splits(run_dir):
    cands = []
    ts = os.path.join(run_dir, "train_summary.json")
    if os.path.isfile(ts):
        try:
            js = json.load(open(ts))
            cands += [(js.get("args") or {}).get("splits_dir"), js.get("splits_dir")]
        except Exception:
            pass
    for dp, _dn, fns in os.walk(run_dir):
        low = os.path.basename(dp).lower()
        for fn in fns:
            if "summary" not in fn.lower() or not fn.endswith(".json"):
                continue
            try:
                js = json.load(open(os.path.join(dp, fn)))
            except Exception:
                continue
            cands.append(js.get("train_splits_dir"))
            if "within" in low or dp == run_dir:
                cands += [js.get("splits_dir"), js.get("eval_splits_dir")]
    for c in cands:
        if c and os.path.isdir(c) and os.path.isfile(os.path.join(c, "train.csv")):
            return c
    return None


def cmd_embed(a):
    """

    Protocol, which must be preserved for the result to mean anything:
      - each fold is fitted on its own training subjects and evaluated on its
        OWN held-out test subjects, so there is no leakage
      - features are never pooled across folds: each fold is a different space
      - the effect size quoted is the gain in mean absolute error IN YEARS over
        predicting the training mean, because R2 is noisy at these fold sizes
      - the probe training set is also capped at common sizes, so probe-size is
        excluded as an explanation for any difference between arms
    """

    torch = _torch()
    from sklearn.linear_model import RidgeCV, LogisticRegression
    from sklearn.metrics import r2_score, mean_absolute_error, roc_auc_score
    from sklearn.preprocessing import StandardScaler

    runs = []
    for root in a.roots:
        if not os.path.isdir(root):
            continue
        for dp, _dn, fns in os.walk(root):
            if "best.pt" in fns:
                runs.append(dp)
    runs = sorted(set(runs))
    if getattr(a, "embed_filter", None):
        runs = [r for r in runs if re.search(a.embed_filter, r.replace("\\", "/"))]
    if not runs:
        die("No best.pt found under {} (filter {})".format(a.roots, a.embed_filter))
    say("Found {} checkpoints.\n".format(len(runs)))

    dev = torch.device("cuda" if (torch.cuda.is_available() and not a.cpu) else "cpu")
    sizes = [int(x) for x in a.probe_sizes.split(",")]
    rows = []

    for rd in runs:
        splits = _find_train_splits(rd)
        if not splits:
            warn("{}: no usable splits_dir, skipping".format(rd)); continue
        tr_csv = os.path.join(splits, "train.csv")
        te_csv = os.path.join(splits, "test.csv")
        if not (os.path.isfile(tr_csv) and os.path.isfile(te_csv)):
            warn("{}: train/test csv missing".format(splits)); continue

        info = parse_cell(os.path.join(rd, "eval_within", "predictions_test.csv"))
        try:
            model, cfg, nc, bk = build_from_checkpoint(a, os.path.join(rd, "best.pt"), torch)
        except SystemExit:
            raise
        except Exception as e:
            warn("{}: {}".format(rd, e)); continue
        model.to(dev)
        run = penultimate_extractor(model, torch)
        patch = a.patch_size
        _ts = os.path.join(rd, "train_summary.json")
        if os.path.isfile(_ts):
            try:
                patch = int((json.load(open(_ts)).get("args") or {})
                            .get("patch_size", a.patch_size))
            except Exception:
                pass

        def feats_of(csv):
            d = pd.read_csv(csv)
            pc, lc2, ac2 = pick_col(d, PATH_PREFS), pick_col(d, LABEL_PREFS), pick_col(d, AGE_PREFS)
            if pc is None:
                return None, None, None
            F, Y, A = [], [], []
            for _i, r in d.iterrows():
                try:
                    v = _load_volume(r[pc], patch)
                except Exception:
                    continue
                x = torch.from_numpy(v)[None, None].float().to(dev)
                _o, f = run(x)
                F.append(f.cpu().numpy().ravel()); Y.append(int(r[lc2])); A.append(float(r[ac2]))
            return np.array(F), np.array(Y), np.array(A)

        say("{}  ({}, {})".format(os.path.basename(rd), bk, "%dcls" % nc))
        Ftr, Ytr, Atr = feats_of(tr_csv)
        Fte, Yte, Ate = feats_of(te_csv)
        if Ftr is None or Fte is None or len(Ftr) < 10 or len(Fte) < 8:
            warn("  too few subjects, skipping"); continue

        cond = balancing_condition(tr_csv, parse_edges(a.age_edges))
        # An adversarial run uses the == splits as its baseline, so the split
        # inspection alone would label it "none" and pool it with the baseline.
        # The distinguishing fact is in the training arguments.
        _tsj = os.path.join(rd, "train_summary.json")
        if os.path.isfile(_tsj):
            try:
                if (json.load(open(_tsj)).get("args") or {}).get("adv_age"):
                    cond = "adversarial" if cond == "none" else cond + "+adversarial"
            except Exception:
                pass
        bal = (cond in ("undersample", "reweight"))
        sc = StandardScaler().fit(Ftr)
        Xtr, Xte = sc.transform(Ftr), sc.transform(Fte)

        base_mae = mean_absolute_error(Ate, np.full(len(Ate), Atr.mean()))
        rr = RidgeCV(alphas=np.logspace(-2, 4, 25)).fit(Xtr, Atr)
        pa = rr.predict(Xte)
        age_r2, mae = r2_score(Ate, pa), mean_absolute_error(Ate, pa)

        dis_auc = np.nan
        try:
            if len(np.unique(Ytr)) > 1 and len(np.unique(Yte)) > 1:
                lg = LogisticRegression(max_iter=2000, C=1.0).fit(Xtr, Ytr)
                P = lg.predict_proba(Xte)
                dis_auc = (roc_auc_score(Yte, P[:, 1]) if P.shape[1] == 2
                           else roc_auc_score(Yte, P, multi_class="ovr", average="macro"))
        except Exception:
            pass

        rec = dict(run=os.path.basename(rd), arm=info["arm"], backbone=bk,
                   cohort=info["train_cohort"], fold=info["fold"],
                   age_bin_balanced=bal, balancing=cond,
                   n_train=len(Xtr), n_test=len(Xte),
                   age_r2=age_r2, mae_model=mae, mae_baseline=base_mae,
                   mae_gain=base_mae - mae, disease_auc=dis_auc, probe_n=len(Xtr))
        # dose-response: cap the probe training set at each common size
        rng = np.random.default_rng(SEED)
        for k in sizes:
            if k >= len(Xtr):
                # capping at or above the available size is just the full fit
                rec["age_r2_n%d" % k] = age_r2; continue
            idx = rng.choice(len(Xtr), k, replace=False)
            r2k = r2_score(Ate, RidgeCV(alphas=np.logspace(-2, 4, 25))
                           .fit(Xtr[idx], Atr[idx]).predict(Xte))
            rec["age_r2_n%d" % k] = r2k
        rows.append(rec)
        say("  age R2 %+.3f | MAE gain %+.2f yr | disease AUC %.3f | balancing: %s"
            % (age_r2, base_mae - mae,
               dis_auc if dis_auc == dis_auc else float("nan"), cond.upper()))
        say("     splits: %s" % splits)

    if not rows:
        die("No probe results produced.")
    df = pd.DataFrame(rows)
    write(df, a.out, "age_probe.csv", a.force)
    cell = df.groupby(["arm", "backbone", "cohort", "balancing"],
                      dropna=False).agg(
        n_folds=("fold", "nunique"), age_r2=("age_r2", "mean"),
        mae_gain=("mae_gain", "mean"), disease_auc=("disease_auc", "mean"),
        **{("age_r2_n%d" % k): ("age_r2_n%d" % k, "mean") for k in sizes}).reset_index()
    write(cell, a.out, "age_probe_cell.csv", a.force)
    say("\nAge decodability per cell:")
    show(cell)
    _probe_permutation(cell, sizes, df)
    return cell


def _probe_permutation(cell, sizes, per_fold=None):
    """Compare age decodability across balancing conditions."""
    if "balancing" not in cell.columns:
        return
    say("\nAge decodability by balancing condition:")
    for c, g in cell.groupby("balancing"):
        say("  {:12s} n={}  age R2 {:+.3f}  MAE gain {:+.2f} yr".format(
            str(c), len(g), g.age_r2.mean(), g.mae_gain.mean()))

    conds = sorted(set(cell.balancing))
    if len(conds) < 2:
        say("\nOnly one balancing condition present, so no contrast is possible.")
        say("To build the ablation you need the SAME cohort, arm and architecture")
        say("under at least two of: none, reweight, undersample.")
        return

    none_r2 = cell[cell.balancing == "none"].age_r2
    for c in conds:
        if c in ("none", "unknown"):
            continue
        g = cell[cell.balancing == c].age_r2
        if not len(none_r2) or not len(g):
            continue
        d = g.mean() - none_r2.mean()
        say("\n  {} versus none: age R2 {:+.3f} -> {:+.3f}  ({:+.3f})".format(
            c, none_r2.mean(), g.mean(), d))
        if d > -0.05:
            say("    This intervention did NOT reduce age decodability.")
            if c == "reweight":
                say("    That is expected on reflection: reweighting keeps every")
                say("    subject, so the age-diagnosis association is still present")
                say("    in every batch and only the gradient magnitude changes.")
                say("    Undersampling is the intervention that removes it.")
        else:
            say("    Age decodability fell under this intervention.")

    # ---- fold-paired comparison, the statistic to quote in the write-up -----
    # Folds are matched by construction: fold k holds the same subjects in every
    # condition. So the conditions can be compared fold by fold rather than as
    # independent groups, which is both correct and far more powerful at n=5.
    if per_fold is not None and "fold" in per_fold.columns:
        from scipy import stats as _st
        # Collapse to exactly one value per fold per condition first. Two runs
        # can legitimately share a fold number under the same condition, and
        # indexing then yields a Series rather than a scalar.
        def _one_per_fold(sub, label):
            g = sub.groupby("fold").age_r2
            if (g.size() > 1).any():
                dup = list(g.size()[g.size() > 1].index)
                warn("{}: fold(s) {} appear more than once; averaging them. "
                     "Check these are not different experiments sharing a "
                     "fold number.".format(label, dup))
            return g.mean()

        base = _one_per_fold(per_fold[per_fold.balancing == "none"], "none")
        for c in sorted(set(per_fold.balancing)):
            if c in ("none", "unknown"):
                continue
            other = _one_per_fold(per_fold[per_fold.balancing == c], c)
            common = sorted(set(base.index) & set(other.index))
            if len(common) < 2:
                say("\n  {} versus none: only {} matched fold(s); train the rest "
                    "before quoting a paired result.".format(c, len(common)))
                continue
            d = np.array([float(other[k]) - float(base[k]) for k in common],
                         dtype=float)
            say("\n  PAIRED across {} folds, {} versus none:".format(len(common), c))
            say("    per-fold change in age R2: {}".format(
                ", ".join("{:+.3f}".format(x) for x in d)))
            say("    mean {:+.3f}   folds reduced: {} of {}".format(
                d.mean(), int((d < 0).sum()), len(d)))
            say("    exact sign test p = {:.4f}".format(sign_test(d)))
            if len(d) >= 5:
                try:
                    w = _st.wilcoxon(d, alternative="two-sided", method="exact")
                    say("    exact Wilcoxon signed-rank p = {:.4f}".format(w.pvalue))
                    say("    (0.0625 is the floor at n=5, so quote it as the minimum"
                        " attainable)")
                except Exception:
                    pass

    n = len(cell)
    m = int((cell.balancing == "none").sum())
    if m and m < n:
        top = set(cell.age_r2.rank(ascending=False).nsmallest(m).index)
        if top == set(cell[cell.balancing == "none"].index):
            say("\nPerfect rank separation: every unbalanced cell outranks every")
            say("balanced one. Exact permutation probability 1 in {} (p = {:.4f})"
                .format(comb(n, m), 1.0 / comb(n, m)))
        else:
            say("\nNo perfect rank separation across conditions.")
    for k in sizes:
        c = "age_r2_n%d" % k
        if c in cell.columns and cell[c].notna().any():
            say("  at probe n={}: {}".format(
                k, {str(kk): round(float(vv), 3)
                    for kk, vv in cell.groupby("balancing")[c].mean().items()}))


# ---------------------------------------------------------------- reorient
def cmd_reorient(a):
    """Summarise a reorientation run and test it fold-wise."""
    from scipy import stats
    p = a.reorient_csv
    if not os.path.isfile(p):
        die("{} not found. Point --reorient-csv at your reorientation output.".format(p))
    d = pd.read_csv(p)
    say("\nReorientation, {} folds.".format(len(d)))
    pairs = [("baseline_auc", "reoriented_auc", "ROC-AUC pooled"),
             ("baseline_auc_OAS1", "reoriented_auc_OAS1", "ROC-AUC OASIS-1"),
             ("baseline_auc_OAS2", "reoriented_auc_OAS2", "ROC-AUC OASIS-2"),
             ("baseline_macro_f1", "reoriented_macro_f1", "macro-F1"),
             ("baseline_bal_acc", "reoriented_bal_acc", "balanced accuracy")]
    rows = []
    for b, r, nm in pairs:
        if b not in d.columns or r not in d.columns:
            continue
        delta = d[r] - d[b]
        try:
            w = stats.wilcoxon(delta, alternative="two-sided", method="exact").pvalue
        except Exception:
            w = np.nan
        rows.append(dict(metric=nm, baseline=d[b].mean(), reoriented=d[r].mean(),
                         change=delta.mean(), folds_improved=int((delta > 0).sum()),
                         n_folds=len(d), sign_test_p=sign_test(delta.values),
                         wilcoxon_p=w,
                         sd_before=d[b].std(ddof=1), sd_after=d[r].std(ddof=1)))
    out = pd.DataFrame(rows)
    write(out, a.out, "reorientation_summary.csv", a.force)
    show(out)
    if "baseline_degenerate" in d.columns:
        say("\nDegenerate folds: {} before, {} after.".format(
            int(d.baseline_degenerate.sum()), int(d.reoriented_degenerate.sum())))
    if {"baseline_auc_OAS1", "baseline_auc_OAS2"} <= set(d.columns):
        gb = d.baseline_auc_OAS1.mean() - d.baseline_auc_OAS2.mean()
        ga = d.reoriented_auc_OAS1.mean() - d.reoriented_auc_OAS2.mean()
        say("Domain gap (OASIS-1 minus OASIS-2): {:.3f} -> {:.3f}".format(gb, ga))
    say("\nAt five folds the smallest attainable exact Wilcoxon p is 0.0625, so quote it")
    say("as 'all five folds improved, p = 0.0625, the minimum attainable at n=5'.")
    return out


# ---------------------------------------------------------------- roi
def cmd_roi(a):
    """Render the group-mean template with the candidate ROI overlaid."""
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt, matplotlib.patches as patches
    try:
        import nibabel as nib
    except ImportError:
        die("nibabel required:  pip install nibabel")
    vols = sorted(glob.glob(os.path.join(a.template_dir, "*.nii*")))
    if not vols:
        vols = sorted(glob.glob(os.path.join(a.template_dir, "**", "*.nii*"),
                                recursive=True))
    if not vols:
        say("No .nii/.nii.gz under {}".format(os.path.abspath(a.template_dir)))
        if os.path.isdir(a.template_dir):
            ext = defaultdict(int)
            for dp, _dn, fns in os.walk(a.template_dir):
                for f in fns:
                    ext[os.path.splitext(f)[1].lower()] += 1
            say("Extensions actually present: {}".format(dict(ext) or "(none)"))
            say("Subdirectories: {}".format(
                sorted(os.listdir(a.template_dir))[:12]))
        else:
            say("That directory does not exist.")
        say("\nTip: point --template-dir at whichever folder holds the")
        say("preprocessed OASIS-1 volumes, or read the scan_path column of a")
        say("split CSV to see where they actually live:")
        say("  python -c \"import pandas as pd;print(pd.read_csv(r'<split>/test.csv').iloc[0])\"")
        sys.exit(1)
    vols = vols[:a.n_template]
    say("Averaging {} volumes...".format(len(vols)))
    acc = None
    for v in vols:
        arr = np.asanyarray(nib.load(v).dataobj).astype(np.float32)
        if acc is None:
            acc = np.zeros_like(arr)
        elif arr.shape != acc.shape:
            continue
        acc += arr
    mean = acc / len(vols)
    box = [int(x) for x in a.roi.split(",")]
    if len(box) != 6:
        die("--roi needs z0,z1,y0,y1,x0,x1")
    z0, z1, y0, y1, x0, x1 = box
    frac = ((z1 - z0) * (y1 - y0) * (x1 - x0)) / float(np.prod(mean.shape))
    cz, cy, cx = (z0 + z1) // 2, (y0 + y1) // 2, (x0 + x1) // 2
    fig, axes = plt.subplots(1, 3, figsize=(13, 4.5))
    for ax, (img, rect, t) in zip(axes, [
            (mean[cz], (x0, y0, x1 - x0, y1 - y0), "axis0 @ {}".format(cz)),
            (mean[:, cy], (x0, z0, x1 - x0, z1 - z0), "axis1 @ {}".format(cy)),
            (mean[:, :, cx], (y0, z0, y1 - y0, z1 - z0), "axis2 @ {}".format(cx))]):
        ax.imshow(img.T, cmap="gray", origin="lower")
        ax.add_patch(patches.Rectangle(rect[:2], rect[2], rect[3], fill=False,
                                       lw=2, edgecolor="red"))
        ax.set_title(t, fontsize=9); ax.axis("off")
    os.makedirs(a.out, exist_ok=True)
    png = os.path.join(a.out, "roi_check.png")
    fig.suptitle("Candidate ROI {} - {:.1f}% of volume".format(box, 100 * frac))
    fig.tight_layout(); fig.savefig(png, dpi=140); plt.close(fig)
    say("\nWritten: {}".format(png))
    say("ROI is {:.2f}% of the volume, so attention ratio 1.0 = exactly what its size "
        "predicts.".format(100 * frac))
    say("\nOPEN IT AND CONFIRM the box sits over medial temporal structures BEFORE running")
    say("attribution. This is the step that failed before: atlas coordinates landed on")
    say("basal ganglia and insula because the processed affine describes a centred grid,")
    say("not an anatomically anchored one.")


# =========================================================================
# ATTRIBUTION: HiResCAM / Grad-CAM / Grad-CAM++ with publication figures
# =========================================================================

def _load_user_module(project_root, needle="def build_model", explicit=None):
    """Import the user's model definitions by finding the file that defines them."""
    import importlib.util
    found = glob.glob(os.path.join(project_root, "**", "*.py"), recursive=True)

    def _rank(f):
        base = os.path.basename(f).lower()
        depth = f.replace("\\", "/").count("/")
        # the module the trainer and evaluator actually import comes first
        pref = 0 if base in ("resnet3d_oasis.py", "medicalnet3d.py",
                             "mobilenet3d_v2.py") else 1
        return (pref, depth, base)

    cands = ([explicit] if explicit else []) + sorted(found, key=_rank)
    skip = (".venv", "site-packages", "__pycache__")
    for c in cands:
        if not c or not os.path.isfile(c) or any(s in c for s in skip):
            continue
        try:
            txt = open(c, encoding="utf-8", errors="ignore").read()
        except Exception:
            continue
        if needle in txt:
            spec = importlib.util.spec_from_file_location("usermodels", c)
            m = importlib.util.module_from_spec(spec)
            sys.path.insert(0, os.path.dirname(os.path.abspath(c)))
            try:
                spec.loader.exec_module(m)
                say("Model definitions loaded from {}".format(c))
                return m, c
            except Exception as e:
                warn("could not import {}: {}".format(c, e))
    return None, None


def _pick_target_layer(model, name=None):
    """
    The layer to attribute from: the last convolutional block before pooling.
    Passing --cam-layer overrides the search.
    """
    import torch.nn as nn
    if name:
        for n, mod in model.named_modules():
            if n == name:
                return n, mod
        die("--cam-layer '{}' not found. Available (last 25):\n  {}".format(
            name, "\n  ".join([n for n, _ in model.named_modules()][-25:])))
    for pref in ("layer4", "features", "blocks", "layer3"):
        for n, mod in model.named_modules():
            if n == pref:
                return n, mod
    last = None
    for n, mod in model.named_modules():
        if isinstance(mod, (nn.Conv3d,)):
            last = (n, mod)
    if last is None:
        die("No Conv3d found. Pass --cam-layer explicitly.")
    return last


def _centre_crop(vol, size):
    out = vol
    for ax in range(3):
        n = out.shape[ax]
        if n <= size:
            continue
        s = (n - size) // 2
        out = np.take(out, range(s, s + size), axis=ax)
    return out


def _load_volume(path, crop=96):
    """
    Load a preprocessed volume and centre-crop it, matching evaluation-time
    behaviour (deterministic centre crop, not the random training crop).
    """
    import nibabel as nib
    arr = np.asanyarray(nib.load(path).dataobj).astype(np.float32)
    return _centre_crop(arr, crop)


def _cam_maps(model, x, target_layer, cls, methods=("hirescam", "gradcam", "gradcampp")):
    """
    Return {method: 3D map} for one input. x is (1,1,D,H,W) and requires_grad.

    HiResCAM is the elementwise product of activations and gradients summed over
    channels (Draelos and Carin, 2020): provably highlights only locations that
    contributed, for a network ending in one fully connected layer. Grad-CAM
    averages the gradient per channel first, which is why it can highlight
    locations the model did not use.
    """
    import torch
    acts, grads = {}, {}

    def fwd(_m, _i, o):
        acts["v"] = o
        o.register_hook(lambda g: grads.__setitem__("v", g))

    h = target_layer.register_forward_hook(fwd)
    model.zero_grad(set_to_none=True)
    out = model(x)
    score = out[0, cls]
    score.backward()
    h.remove()

    A = acts["v"].detach()[0]          # (C, d, h, w)
    G = grads["v"].detach()[0]
    maps = {}
    if "hirescam" in methods:
        maps["hirescam"] = torch.relu((A * G).sum(0))
    if "gradcam" in methods:
        w = G.mean(dim=(1, 2, 3), keepdim=True)
        maps["gradcam"] = torch.relu((w * A).sum(0))
    if "gradcampp" in methods:
        g2, g3 = G ** 2, G ** 3
        denom = 2 * g2 + (A * g3).sum(dim=(1, 2, 3), keepdim=True)
        alpha = g2 / torch.clamp(denom, min=1e-8)
        w = (alpha * torch.relu(G)).sum(dim=(1, 2, 3), keepdim=True)
        maps["gradcampp"] = torch.relu((w * A).sum(0))
    return {k: v.cpu().numpy() for k, v in maps.items()}, float(out[0, cls]), int(out.argmax())


def _upsample(m, shape):
    from scipy.ndimage import zoom
    f = [s / c for s, c in zip(shape, m.shape)]
    return zoom(m, f, order=1)


def _norm(m):
    m = np.asarray(m, float)
    lo, hi = np.nanpercentile(m, 1), np.nanpercentile(m, 99)
    return np.clip((m - lo) / (hi - lo + 1e-9), 0, 1)


def _montage(ax, bg, cam, axis, n_slices, title, cmap="inferno", alpha=0.45,
             thresh=0.35):
    """One row of the figure: n_slices evenly spaced slices along `axis`."""
    D = bg.shape[axis]
    idx = np.linspace(int(0.25 * D), int(0.75 * D), n_slices).astype(int)
    tiles_bg, tiles_cam = [], []
    for i in idx:
        b = np.take(bg, i, axis=axis)
        c = np.take(cam, i, axis=axis)
        tiles_bg.append(np.rot90(b)); tiles_cam.append(np.rot90(c))
    B = np.concatenate(tiles_bg, axis=1)
    C = np.concatenate(tiles_cam, axis=1)
    ax.imshow(B, cmap="gray",
              vmin=np.percentile(B, 1), vmax=np.percentile(B, 99.5))
    M = np.ma.masked_where(C < thresh, C)
    im = ax.imshow(M, cmap=cmap, alpha=alpha, vmin=0, vmax=1)
    ax.set_ylabel(title, fontsize=9)
    ax.set_xticks([]); ax.set_yticks([])
    return im


def build_from_checkpoint(a, ckpt_path, torch):
    """
    Rebuild the exact architecture a checkpoint came from, the way the
    evaluator does: the saved cfg identifies the family by which fields it has.
        MedicalNetConfig    -> model_depth
        MobileNet3DV2Config -> width_mult
        ModelConfig         -> backbone
    Returns (model, cfg_dict, num_classes, backbone_name).
    """
    ck = torch.load(ckpt_path, map_location="cpu", weights_only=False)
    cfg = ck.get("cfg", {}) or {}
    if not isinstance(cfg, dict):
        try:
            from dataclasses import asdict as _ad
            cfg = _ad(cfg)
        except Exception:
            cfg = {}
    nc = int(ck.get("num_classes") or cfg.get("out_dim") or 2)

    def _mod(needle):
        m, _ = _load_user_module(a.project_root, needle, a.model_module)
        return m

    model = name = None
    if "model_depth" in cfg:
        m = _mod("def build_medicalnet")
        if m:
            MC, bm = getattr(m, "MedicalNetConfig"), getattr(m, "build_medicalnet")
            f = getattr(MC, "__dataclass_fields__", {})
            kw = {k: v for k, v in cfg.items() if k in f}
            kw.update(out_dim=nc)
            if "pretrained_path" in f:
                kw["pretrained_path"] = None
            if "no_cuda" in f:
                # MedicalNet's downsample_basic_block allocates its zero-pad
                # tensor on CPU and only moves it to CUDA when no_cuda is False.
                # Forcing True put the pad on CPU under a CUDA model.
                kw["no_cuda"] = bool(getattr(a, "cpu", False)) or not torch.cuda.is_available()
            model, name = bm(MC(**kw)), "medicalnet_resnet18"
    elif "width_mult" in cfg:
        m = _mod("def build_mobilenet3d_v2")
        if m:
            MC, bm = getattr(m, "MobileNet3DV2Config"), getattr(m, "build_mobilenet3d_v2")
            f = getattr(MC, "__dataclass_fields__", {})
            kw = {k: v for k, v in cfg.items() if k in f}
            kw.update(out_dim=nc)
            if "pretrained_path" in f:
                kw["pretrained_path"] = None
            model, name = bm(MC(**kw)), "mobilenet3d_v2"
    if model is None:
        m = _mod("def build_model")
        if m is None:
            die("Could not find build_model. Pass --model-module PATH.")
        MC = getattr(m, "ModelConfig", None) or getattr(m, "ModelCfg", None)
        bm = getattr(m, "build_model")
        bk = cfg.get("backbone", a.backbone)
        if MC is not None:
            f = getattr(MC, "__dataclass_fields__", {})
            kw = {k: v for k, v in cfg.items() if k in f}
            kw.setdefault("backbone", bk); kw.update(out_dim=nc)
            if "pretrained" in f:
                kw["pretrained"] = False
            if "in_channels" in f:
                kw.setdefault("in_channels", 1)
            model = bm(MC(**kw))
        else:
            model = bm(backbone=bk, out_dim=nc, in_channels=1, pretrained=False)
        name = str(bk)

    sd = ck.get("model_state") or ck.get("state_dict") or ck.get("model")
    if sd is None:
        die("Checkpoint has no weights. Keys: {}".format(list(ck.keys())))
    res = model.load_state_dict(sd, strict=False)
    nm = len(getattr(res, "missing_keys", []))
    if nm > 10:
        warn("{} missing keys rebuilding {}; architecture may not match."
             .format(nm, os.path.basename(ckpt_path)))
    model.eval()
    return model, cfg, nc, name


def penultimate_extractor(model, torch):
    """Return a function x -> (logits, penultimate features)."""
    import torch.nn as nn
    head = None
    for _n, m in model.named_modules():
        if isinstance(m, nn.Linear):
            head = m
    if head is None:
        die("No Linear head found; cannot extract penultimate features.")
    store = {}
    head.register_forward_hook(lambda _m, inp, _o: store.__setitem__("v", inp[0].detach()))

    def run(x):
        with torch.no_grad():
            out = model(x)
        return out, store["v"]
    return run


def balancing_condition(train_csv, edges=None):
    """
    Name the balancing intervention a training split received.

    'undersample' and 'reweight' are NOT the same intervention and must never
    be pooled. Undersampling removes subjects, so the network never sees the
    age-diagnosis association. Reweighting keeps every subject and only scales
    the loss, so the association is still present in every batch. Reporting
    them under one boolean hides exactly the contrast that matters.
    """
    try:
        d = pd.read_csv(train_csv)
    except Exception:
        return "unknown"
    ac, lc = pick_col(d, AGE_PREFS), pick_col(d, LABEL_PREFS)
    if ac is None or lc is None:
        return "unknown"
    wc = next((c for c in d.columns
               if str(c).lower() in ("sample_weight", "weight", "sw")), None)
    d = d.copy()
    d["_b"] = assign_bins(d[ac].values, edges or AGE_EDGES)
    d["_one"] = 1.0

    def _uniform(col):
        for _b, g in d.groupby("_b"):
            if len(g) < 3:
                continue
            mass = g.groupby(lc)[col].sum()
            if len(mass) > 1 and (mass.max() - mass.min()) > max(1e-6, 0.15 * mass.max()):
                return False
        return True

    if _uniform("_one"):
        return "undersample"
    if wc is not None:
        d["_w"] = d[wc].astype(float)
        if _uniform("_w"):
            return "reweight"
    return "none"


def is_age_bin_balanced(train_csv, edges=None):
    """
    Measure whether the EFFECTIVE class prior is uniform within each age bin.

    This must account for both balancing modes, because they look completely
    different on disk:
      undersample - the row COUNTS are equalised within each bin
      reweight    - the counts are untouched and a sample_weight column carries
                    the correction, so counting rows would wrongly report the
                    arm as unbalanced

    So the check sums sample_weight per class per bin when that column exists,
    and falls back to counting rows when it does not.
    Returns True, False, or None when it cannot be determined.
    """
    try:
        d = pd.read_csv(train_csv)
    except Exception:
        return None
    ac, lc = pick_col(d, AGE_PREFS), pick_col(d, LABEL_PREFS)
    if ac is None or lc is None:
        return None
    wc = next((c for c in d.columns
               if str(c).lower() in ("sample_weight", "weight", "sw")), None)
    d = d.copy()
    d["_b"] = assign_bins(d[ac].values, edges or AGE_EDGES)
    d["_w"] = d[wc].astype(float) if wc else 1.0

    def _uniform(col):
        """Is the per-class mass in `col` uniform inside every populated bin?"""
        for _b, g in d.groupby("_b"):
            if len(g) < 3:
                continue
            mass = g.groupby(lc)[col].sum()
            if len(mass) > 1 and (mass.max() - mass.min()) > max(1e-6, 0.15 * mass.max()):
                return False
        return True

    d["_one"] = 1.0
    # Either route counts as balanced: equalised row counts (undersample) or
    # equalised weighted mass (reweight). Checking both also survives a stale
    # or irrelevant weight column sitting beside already-equal counts.
    return bool(_uniform("_one") or (wc is not None and _uniform("_w")))


def cmd_cam(a):
    """
    Volumetric attribution with publication-quality figures.

    Produces, per requested method:
      - a three-plane multi-slice montage of the GROUP MEAN map
      - the same split by diagnosis and by age bin
      - a per-subject sheet for the clearest examples
      - the region-of-interest attention ratio
      - the Adebayo randomisation sanity check
    """
    torch = _torch()
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    try:
        import nibabel  # noqa: F401
    except ImportError:
        die("nibabel required:  pip install nibabel")

    if not os.path.isfile(a.checkpoint):
        die("--checkpoint not found: {}".format(a.checkpoint))
    split = a.cam_split or os.path.join(a.splits_dir, "test.csv")
    if not os.path.isfile(split):
        die("Split CSV not found: {}. Pass --cam-split.".format(split))

    df = pd.read_csv(split)
    pc = pick_col(df, PATH_PREFS)
    lc = pick_col(df, LABEL_PREFS)
    ac = pick_col(df, AGE_PREFS)
    sc = pick_col(df, SUBJ_PREFS)
    if pc is None or lc is None:
        die("Split CSV needs a scan-path and a label column. Found: {}".format(list(df.columns)))

    # Group-level maps require a shared anatomical space. Only the OASIS-1
    # subset has one: its volumes share a bounding box to within 1.7 per cent.
    if a.oas1_only and sc:
        keep = df[sc].astype(str).str.contains(a.oas1_pattern, case=False, regex=True)
        if keep.sum() == 0:
            warn("No subject matched --oas1-pattern; using all rows.")
        else:
            say("Restricting to {} OASIS-1 subjects (shared anatomical space).".format(int(keep.sum())))
            df = df[keep].reset_index(drop=True)
    if a.cam_n and len(df) > a.cam_n:
        # Stratified subsample by index. groupby(...).apply(...) is NOT used here:
        # in pandas 2.2+ it excludes the grouping column from the frames it hands
        # the callback, so the label column vanished and every row lookup failed.
        per = max(1, a.cam_n // max(1, df[lc].nunique()))
        take = []
        for _v, g in df.groupby(lc, sort=True):
            take.extend(g.sample(min(len(g), per), random_state=SEED).index.tolist())
        df = df.loc[sorted(take)].reset_index(drop=True)
        assert lc in df.columns, "label column lost during subsampling"
    say("Attributing over {} subjects.".format(len(df)))

    # Build the model exactly the way the trainer does. This project uses
    #   from resnet3d_OASIS import ModelConfig, build_model      -> build_model(cfg)
    #   from medicalnet3d   import MedicalNetConfig, build_medicalnet
    # so build_model takes a config OBJECT, not keyword arguments.
    ck = torch.load(a.checkpoint, map_location="cpu", weights_only=False)
    cfg = ck.get("cfg", {}) or {}
    if not isinstance(cfg, dict):
        try:
            from dataclasses import asdict as _asdict
            cfg = _asdict(cfg)
        except Exception:
            cfg = {}
    nc = ck.get("num_classes") or cfg.get("out_dim") or 2
    bk = cfg.get("backbone", a.backbone)

    model = None
    if str(bk) == "medicalnet_resnet18":
        mod, _src = _load_user_module(a.project_root, "def build_medicalnet",
                                      a.model_module)
        if mod is not None:
            MC = getattr(mod, "MedicalNetConfig", None)
            bm = getattr(mod, "build_medicalnet", None)
            if MC and bm:
                try:
                    model = bm(MC(**{k: v for k, v in cfg.items()
                                     if k in getattr(MC, "__dataclass_fields__", {})}))
                except Exception as e:
                    warn("build_medicalnet failed: {}".format(e))
    if model is None:
        mod, _src = _load_user_module(a.project_root, "def build_model", a.model_module)
        if mod is None:
            die("Could not find a module defining build_model. Pass --model-module PATH.")
        MC = getattr(mod, "ModelConfig", None) or getattr(mod, "ModelCfg", None)
        bmf = getattr(mod, "build_model", None)
        if MC is not None:
            fields = getattr(MC, "__dataclass_fields__", {})
            kw = {k: v for k, v in cfg.items() if k in fields}
            kw.setdefault("backbone", bk)
            kw.setdefault("out_dim", nc)
            if "in_channels" in fields:
                kw.setdefault("in_channels", 1)
            if "pretrained" in fields:
                kw["pretrained"] = False       # weights come from the checkpoint
            if "dropout" in fields:
                kw.setdefault("dropout", cfg.get("dropout", 0.2))
            model = bmf(MC(**kw))
        else:
            try:
                model = bmf(backbone=bk, out_dim=nc, in_channels=1,
                            pretrained=False, dropout=cfg.get("dropout", 0.2))
            except TypeError:
                model = bmf(bk, nc)

    sd = (ck.get("model_state") or ck.get("state_dict") or ck.get("model"))
    if sd is None:
        die("Checkpoint has no weights under model_state/state_dict/model. "
            "Keys present: {}".format(list(ck.keys())))
    res = model.load_state_dict(sd, strict=False)
    n_missing = len(getattr(res, "missing_keys", []))
    n_unexp = len(getattr(res, "unexpected_keys", []))
    say("Checkpoint loaded: {} missing, {} unexpected keys.".format(n_missing, n_unexp))
    if n_missing > 10:
        warn("Many missing keys. The architecture may not match the checkpoint; "
             "attribution would be meaningless. Check --backbone.")
    dev = torch.device("cuda" if (torch.cuda.is_available() and not a.cpu) else "cpu")
    model.to(dev).eval()

    lname, layer = _pick_target_layer(model, a.cam_layer)
    say("Attributing from layer: {}".format(lname))

    methods = [m.strip() for m in a.cam_methods.split(",")]
    acc = {m: None for m in methods}
    by_class, by_bin, n_ok = defaultdict(list), defaultdict(list), 0
    bg_acc = None
    rows, examples = [], []

    for i, r in df.iterrows():
        try:
            vol = _load_volume(r[pc], a.patch_size)
        except Exception as e:
            warn("{}: {}".format(r[pc], e)); continue
        x = torch.from_numpy(vol)[None, None].float().to(dev).requires_grad_(True)
        cls = int(r[lc]) if a.cam_target == "true" else None
        with torch.enable_grad():
            if cls is None:
                with torch.no_grad():
                    cls = int(model(x).argmax())
            maps, score, pred = _cam_maps(model, x, layer, cls, methods)
        big = {m: _norm(_upsample(v, vol.shape)) for m, v in maps.items()}
        for m in methods:
            acc[m] = big[m] if acc[m] is None else acc[m] + big[m]
        bg_acc = vol.astype(np.float64) if bg_acc is None else bg_acc + vol
        by_class[int(r[lc])].append(big[methods[0]])
        if ac:
            b = assign_bins([r[ac]])[0]; by_bin[b].append(big[methods[0]])
        n_ok += 1
        rows.append(dict(subject=r[sc] if sc else i, label=int(r[lc]),
                         age=float(r[ac]) if ac else np.nan, pred=pred, score=score))
        if len(examples) < a.cam_examples:
            examples.append((r[sc] if sc else str(i), vol, big[methods[0]], int(r[lc]), pred))

    if n_ok == 0:
        die("No volumes could be loaded. Check the scan_path column in {}".format(split))
    for m in methods:
        acc[m] /= n_ok
    bg = bg_acc / n_ok
    say("Averaged {} subjects.".format(n_ok))

    os.makedirs(a.out, exist_ok=True)
    figs = []

    # ---- Figure 6: group-mean, three planes, per method ---------------------
    for m in methods:
        fig, axes = plt.subplots(3, 1, figsize=(a.cam_slices * 1.9, 6.4))
        for ax, axis, nm in zip(axes, (0, 1, 2), ("axial", "coronal", "sagittal")):
            im = _montage(ax, bg, acc[m], axis, a.cam_slices, nm,
                          cmap=a.cam_cmap, alpha=a.cam_alpha, thresh=a.cam_thresh)
        cb = fig.colorbar(im, ax=axes, fraction=0.015, pad=0.01)
        cb.set_label("normalised attribution", fontsize=8)
        fig.suptitle("Group-mean {} over {} subjects  (target layer {})".format(
            m.upper(), n_ok, lname), fontsize=11)
        p = os.path.join(a.out, "cam_group_{}.png".format(m))
        fig.savefig(p, dpi=a.cam_dpi, bbox_inches="tight"); plt.close(fig)
        figs.append(p); say("  written: {}".format(p))

    # ---- group mean split by diagnosis --------------------------------------
    if len(by_class) > 1:
        ks = sorted(by_class)
        fig, axes = plt.subplots(len(ks), 1, figsize=(a.cam_slices * 1.9, 2.2 * len(ks)))
        for ax, c in zip(np.atleast_1d(axes), ks):
            g = np.mean(by_class[c], axis=0)
            im = _montage(ax, bg, g, 0, a.cam_slices,
                          "class {} (n={})".format(c, len(by_class[c])),
                          cmap=a.cam_cmap, alpha=a.cam_alpha, thresh=a.cam_thresh)
        fig.colorbar(im, ax=axes, fraction=0.015, pad=0.01)
        fig.suptitle("Group-mean {} by diagnosis (axial)".format(methods[0].upper()), fontsize=11)
        p = os.path.join(a.out, "cam_by_diagnosis.png")
        fig.savefig(p, dpi=a.cam_dpi, bbox_inches="tight"); plt.close(fig)
        figs.append(p); say("  written: {}".format(p))

    # ---- group mean split by age bin ----------------------------------------
    if len(by_bin) > 1:
        ks = [b for b in ("<70", "70-80", "80+") if b in by_bin]
        fig, axes = plt.subplots(len(ks), 1, figsize=(a.cam_slices * 1.9, 2.2 * len(ks)))
        for ax, b in zip(np.atleast_1d(axes), ks):
            g = np.mean(by_bin[b], axis=0)
            im = _montage(ax, bg, g, 0, a.cam_slices, "age {} (n={})".format(b, len(by_bin[b])),
                          cmap=a.cam_cmap, alpha=a.cam_alpha, thresh=a.cam_thresh)
        fig.colorbar(im, ax=axes, fraction=0.015, pad=0.01)
        fig.suptitle("Group-mean {} by age bin (axial)".format(methods[0].upper()), fontsize=11)
        p = os.path.join(a.out, "cam_by_age_bin.png")
        fig.savefig(p, dpi=a.cam_dpi, bbox_inches="tight"); plt.close(fig)
        figs.append(p); say("  written: {}".format(p))

    # ---- method comparison on the group mean --------------------------------
    if len(methods) > 1:
        fig, axes = plt.subplots(len(methods), 1,
                                 figsize=(a.cam_slices * 1.9, 2.2 * len(methods)))
        for ax, m in zip(np.atleast_1d(axes), methods):
            im = _montage(ax, bg, acc[m], 0, a.cam_slices, m.upper(),
                          cmap=a.cam_cmap, alpha=a.cam_alpha, thresh=a.cam_thresh)
        fig.colorbar(im, ax=axes, fraction=0.015, pad=0.01)
        fig.suptitle("Attribution method comparison, group mean (axial)", fontsize=11)
        p = os.path.join(a.out, "cam_method_comparison.png")
        fig.savefig(p, dpi=a.cam_dpi, bbox_inches="tight"); plt.close(fig)
        figs.append(p); say("  written: {}".format(p))

    # ---- per-subject examples ----------------------------------------------
    if examples:
        fig, axes = plt.subplots(len(examples), 1,
                                 figsize=(a.cam_slices * 1.9, 2.2 * len(examples)))
        for ax, (sid, v, c, lab, pr) in zip(np.atleast_1d(axes), examples):
            im = _montage(ax, v, c, 0, a.cam_slices,
                          "{}\ntrue {} pred {}".format(str(sid)[:14], lab, pr),
                          cmap=a.cam_cmap, alpha=a.cam_alpha, thresh=a.cam_thresh)
        fig.colorbar(im, ax=axes, fraction=0.015, pad=0.01)
        fig.suptitle("Individual subjects, {} (axial)".format(methods[0].upper()), fontsize=11)
        p = os.path.join(a.out, "cam_examples.png")
        fig.savefig(p, dpi=a.cam_dpi, bbox_inches="tight"); plt.close(fig)
        figs.append(p); say("  written: {}".format(p))

    # ---- ROI attention ratio ------------------------------------------------
    box = [int(x) for x in a.roi.split(",")]
    if len(box) == 6:
        sc_ = a.patch_size / 128.0
        z0, z1, y0, y1, x0, x1 = [int(round(b * sc_)) for b in box]
        z1 = min(z1, bg.shape[0]); y1 = min(y1, bg.shape[1]); x1 = min(x1, bg.shape[2])
        frac = ((z1 - z0) * (y1 - y0) * (x1 - x0)) / float(np.prod(bg.shape))
        say("\nROI attention ratio (1.0 = exactly the share its volume predicts):")
        for m in methods:
            inside = acc[m][z0:z1, y0:y1, x0:x1].mean()
            ratio = inside / (acc[m].mean() + 1e-9)
            say("  {:12s} {:.3f}   (ROI is {:.1f}% of volume)".format(m, ratio, 100 * frac))
        say("  Confirm the box first:  python msc_analysis.py roi --roi {}".format(a.roi))

    # ---- Adebayo randomisation sanity check ---------------------------------
    if a.cam_sanity:
        say("\nRandomisation sanity check (Adebayo et al., 2018)...")
        import copy
        rnd = copy.deepcopy(model)
        for p_ in rnd.parameters():
            if p_.dim() > 1:
                torch.nn.init.xavier_uniform_(p_)
        _ln, rl = _pick_target_layer(rnd, a.cam_layer)
        rnd.to(dev).eval()
        acc_r, nn_ = None, 0
        for _i, r in df.head(min(len(df), a.cam_sanity_n)).iterrows():
            try:
                vol = _load_volume(r[pc], a.patch_size)
            except Exception:
                continue
            x = torch.from_numpy(vol)[None, None].float().to(dev).requires_grad_(True)
            mm, _s, _p = _cam_maps(rnd, x, rl, int(r[lc]), [methods[0]])
            u = _norm(_upsample(mm[methods[0]], vol.shape))
            acc_r = u if acc_r is None else acc_r + u
            nn_ += 1
        if nn_:
            acc_r /= nn_
            corr = float(np.corrcoef(acc[methods[0]].ravel(), acc_r.ravel())[0, 1])
            say("  correlation(trained, randomised) = {:.3f} over {} subjects".format(corr, nn_))
            say("  Near 0 means the map reflects the LEARNED weights, which is the")
            say("  result you want. Near 1 would mean it reflects architecture alone")
            say("  and explains nothing, and the attribution section should be cut.")

    out = pd.DataFrame(rows)
    write(out, a.out, "cam_subjects.csv", a.force)
    say("\nFigures written: {}".format(len(figs)))
    say("Use cam_group_{}.png as Figure 6.".format(methods[0]))
    return out


# =========================================================================
# GENERATORS
# =========================================================================

def detect_scripts(root="."):
    """
    Find the trainer and the per-checkpoint evaluator.

    Matching on the string "predictions_test.csv" alone is not enough: an
    analysis aggregator that READS those files matches just as well as the
    evaluator that WRITES them. Candidates are therefore scored on the
    arguments they declare. A real evaluator takes a checkpoint and an output
    directory; an aggregator takes --roots and is disqualified.
    """
    skip = {".venv", "venv", "site-packages", "__pycache__", ".git", "node_modules"}
    me = os.path.basename(__file__)
    tr, ev = [], []
    for dp, dn, fns in os.walk(root):
        dn[:] = [d for d in dn if d not in skip]
        for fn in fns:
            if not fn.endswith(".py") or fn == me:
                continue
            full = os.path.join(dp, fn)
            try:
                t = open(full, encoding="utf-8", errors="ignore").read()
            except Exception:
                continue

            # --- evaluator score ---
            sc = 0
            if "predictions_test.csv" in t:
                sc += 1
            if re.search(r'add_argument\(\s*["\']--checkpoint', t):
                sc += 3
            if re.search(r'add_argument\(\s*["\']--out_?dir', t):
                sc += 2
            if re.search(r'predictions_test\.csv["\']?\s*\)?\s*(,|\))?', t) and "to_csv" in t:
                sc += 2
            if re.search(r'add_argument\(\s*["\']--roots', t):
                sc -= 6          # aggregator, not an evaluator
            if "load_state_dict" in t:
                sc += 2
            base = os.path.basename(full).lower()
            if base.startswith("eval") or "evaluate" in base:
                sc += 4                     # named like an evaluator
            for bad in ("reorient", "gradcam", "cam_", "calib", "verify",
                        "figure", "plot", "analysis", "aggregate", "probe"):
                if bad in base:
                    sc -= 5                 # a tool that happens to touch predictions
            if re.search(r"\(\d+\)", base):
                sc -= 2                     # "evaluate (1).py" duplicate copies
            if sc >= 4:
                ev.append((sc, full))

            # --- trainer score ---
            st = 0
            if "train_summary.json" in t:
                st += 3
            if "torch.save" in t:
                st += 2
            if re.search(r'add_argument\(\s*["\']--epochs', t):
                st += 2
            if re.search(r'add_argument\(\s*["\']--results_?dir', t):
                st += 1
            if re.search(r'add_argument\(\s*["\']--roots', t):
                st -= 6
            if st >= 4:
                tr.append((st, full))

    tr.sort(reverse=True); ev.sort(reverse=True)
    return ((tr[0][1] if tr else None), (ev[0][1] if ev else None),
            [f for _s, f in tr], [f for _s, f in ev])


def script_args(path):
    """Return the set of long options a script declares."""
    if not path or not os.path.isfile(path):
        return set()
    try:
        t = open(path, encoding="utf-8", errors="ignore").read()
    except Exception:
        return set()
    return set(re.findall(r"""add_argument\(\s*["']--([A-Za-z0-9_\-]+)""", t))


def eval_flags(path):
    """
    Map this evaluator's real flag names. Scripts differ: yours takes
    --ckpt / --results_dir, not --checkpoint / --out_dir, and guessing wrong
    produces a .bat file where every line fails on 'unrecognized arguments'.
    """
    have = script_args(path)
    def first(cands, fallback):
        for c in cands:
            if c in have:
                return c
        return fallback
    return dict(
        ckpt=first(["checkpoint", "ckpt", "weights", "model_path", "model"], "ckpt"),
        splits=first(["splits_dir", "splits", "data_dir"], "splits_dir"),
        out=first(["out_dir", "outdir", "output_dir", "results_dir", "out"], "results_dir"),
        known=bool(have))


def eval_cmd(script, ckpt, splits, out):
    f = eval_flags(script)
    return 'python "{s}" --{c} "{ck}" --{sp} "{spd}" --{o} "{od}"'.format(
        s=script, c=f["ckpt"], ck=ckpt, sp=f["splits"], spd=splits, o=f["out"], od=out)


def cmd_report(a):
    """
    Assemble every CSV in the output directory into one results document.
    Writes markdown always, and .docx as well when python-docx is available.
    Re-run it after any analysis and the document is current.
    """
    src = a.out
    files = sorted(glob.glob(os.path.join(src, "*.csv")))
    if not files:
        die("No CSVs in {}. Run `all` first.".format(src))

    ORDER = [
        ("run_index.csv", "Every run", "One row per prediction file: arm, backbone, "
         "cohorts, fold, and its headline metrics. The index the rest is built from."),
        ("cv_performance.csv", "Performance for every cell",
         "Within- and cross-cohort fold means with bootstrap intervals. "
         "auc_above_chance=False means the interval contains 0.500."),
        ("degradation_by_metric.csv", "Within-to-cross degradation by metric",
         "macro-F1 change against ROC-AUC change. A ratio above 1 means the "
         "decision rule degraded faster than the ranking."),
        ("paired_backbone.csv", "Paired backbone comparisons",
         "Bootstrap on the difference using identical resamples. Comparing two "
         "separate marginal intervals would be the wrong test."),
        ("per_bin_metrics_foldmean.csv", "Per-age-bin rates, fold means",
         "Sensitivity, specificity and FNR for the disease class by age bin. "
         "Rows with any_degenerate=True are not age effects."),
        ("per_bin_metrics.csv", "Per-age-bin rates, per fold", "The same, unaggregated."),
        ("age_bin_sensitivity.csv", "Age gap under alternative binnings",
         "The worst-minus-best gap under six binning schemes. Stability across "
         "schemes means the gap is not an artefact of the 70/80 cut-points."),
        ("confidence_age_foldmean.csv", "Confidence and posterior versus age",
         "Positive r_age_disease is the quantitative form of age-tracking."),
        ("threshold_interventions.csv", "Threshold interventions",
         "Baseline, prior correction, prevalence matching and the oracle bound. "
         "The oracle uses test labels and is a decomposition, never a method."),
        ("age_gap_paired.csv", "Paired bootstrap on the age gap",
         "The primary success criterion. Negative difference means the gap "
         "narrowed. Cells with degenerate folds are excluded by protocol."),
        ("degeneracy_inventory.csv", "Degeneracy inventory",
         "Every cell and intervention, and how many folds collapsed to a single "
         "predicted class. The audit trail for every exclusion made elsewhere."),
        ("domain_stratified_foldmean.csv", "OASIS-1 versus OASIS-2",
         "The pooled AUC sits below both subsets because cross-subset pairs are "
         "ordered by preprocessing pipeline rather than by disease."),
        ("calibration_by_age_cell.csv", "Calibration by age bin",
         "Binned expected calibration error. A constant predictor at the base "
         "rate is perfectly calibrated and useless, so read with the degeneracy flag."),
        ("dro_results.csv", "Group distributionally robust optimisation",
         "Per-fold held-out test results for the in-processing intervention."),
        ("age_probe.csv", "Age decodability, per fold",
         "Ridge regression predicting chronological age from the penultimate "
         "features, fitted and evaluated within each fold."),
        ("age_probe_cell.csv", "Age decodability, per cell",
         "The same aggregated, with the dose-response at capped probe sizes."),
        ("reorientation_summary.csv", "Reorientation",
         "Effect of aligning volumes to a common voxel axis order at evaluation."),
        ("seed_sensitivity.csv", "Repeated-seed sensitivity",
         "Seed-to-seed variation, to test whether fold variance is a fair stand-in."),
        ("balanced_split_summary.csv", "Balanced split construction",
         "Subjects retained when the class prior is equalised within each age bin."),
        ("checkpoint_config.csv", "Training configuration read from checkpoints",
         "Settings verified by measurement rather than recalled."),
        ("cam_subjects.csv", "Attribution, per subject",
         "Region attention ratio per subject, with true and predicted label."),
    ]
    def _read(path):
        """An analysis that produced no rows leaves an empty file; skip it."""
        try:
            d = pd.read_csv(path)
            return d if len(d.columns) else None
        except Exception as e:
            warn("{}: {}".format(os.path.basename(path), e))
            return None

    seen, blocks = set(), []
    for name, title, note in ORDER:
        p = os.path.join(src, name)
        if os.path.isfile(p):
            d = _read(p); seen.add(name)
            if d is not None:
                blocks.append((name, title, note, d))
    for p in files:
        n = os.path.basename(p)
        if n in seen:
            continue
        d = _read(p)
        if d is not None:
            blocks.append((n, n.replace("_", " ").replace(".csv", "").capitalize(), "", d))
    if not blocks:
        die("Every CSV in {} was empty or unreadable.".format(src))

    os.makedirs(src, exist_ok=True)
    md = os.path.join(src, "ALL_RESULTS.md")
    with open(md, "w", encoding="utf-8") as f:
        f.write("# All results\n\n")
        f.write("Every value here is a held-out test score computed from a saved "
                "per-subject prediction file. No training-time or validation-time "
                "figure appears. Chance is 0.333 for macro-F1 and balanced accuracy "
                "in the three-class arm and 0.500 elsewhere.\n\n")
        f.write("The `variant` column separates baseline runs from the Group DRO "
                "intervention. Filter to `variant == base` for the 24 cells the "
                "main results tables describe.\n\n")
        for i, (n, t, note, d) in enumerate(blocks, 1):
            f.write("\n## Table {}. {}\n\n".format(i, t))
            if note:
                f.write("{}\n\n".format(note))
            f.write("*Source: `{}`, {} rows.*\n\n".format(n, len(d)))
            dd = d.copy()
            for c in dd.columns:
                if dd[c].dtype.kind == "f":
                    dd[c] = dd[c].round(4)
            if len(dd) > a.report_max_rows:
                f.write(dd.head(a.report_max_rows).to_markdown(index=False))
                f.write("\n\n*(first {} of {} rows; full data in the CSV)*\n"
                        .format(a.report_max_rows, len(dd)))
            else:
                f.write(dd.to_markdown(index=False))
            f.write("\n")
    say("written: {}".format(md))

    try:
        from docx import Document
        from docx.shared import Pt, Inches
    except ImportError:
        say("\npython-docx not installed, so only markdown was written.")
        say("  pip install python-docx     then re-run to get the .docx")
        return
    doc = Document()
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Inches(11.69), Inches(8.27)   # landscape A4
    sec.left_margin = sec.right_margin = Inches(0.6)
    sec.top_margin = sec.bottom_margin = Inches(0.6)
    st = doc.styles["Normal"]; st.font.name = "Calibri"; st.font.size = Pt(9)

    def head(txt, size=16, before=0):
        p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(before)
        p.paragraph_format.keep_with_next = True
        r = p.add_run(txt); r.bold = True; r.font.size = Pt(size)

    def para(txt, size=9, italic=False):
        p = doc.add_paragraph(); r = p.add_run(txt)
        r.font.size = Pt(size); r.italic = italic

    head("All results", 20)
    para("Robustness of 3D CNN Alzheimer's Disease Classification to Age "
         "Distribution Shift")
    para("Every value is a held-out test score from a saved per-subject "
         "prediction file. No training-time or validation-time figure appears. "
         "Chance is 0.333 for macro-F1 and balanced accuracy in the three-class "
         "arm and 0.500 elsewhere. The variant column separates baseline runs "
         "from the Group DRO intervention; filter to variant == base for the 24 "
         "cells the main results tables describe.")
    for i, (n, t, note, d) in enumerate(blocks, 1):
        head("Table {}. {}".format(i, t), 12, before=12)
        if note:
            para(note)
        para("Source: {}, {} rows.".format(n, len(d)), italic=True)
        dd = d.head(a.report_max_rows).copy()
        for c in dd.columns:
            if dd[c].dtype.kind == "f":
                dd[c] = dd[c].round(4)
        cols = list(dd.columns)[:a.report_max_cols]
        tab = doc.add_table(rows=len(dd) + 1, cols=len(cols)); tab.style = "Table Grid"
        for ci, c in enumerate(cols):
            cell = tab.cell(0, ci); cell.text = ""
            r = cell.paragraphs[0].add_run(str(c)); r.bold = True; r.font.size = Pt(7.5)
        for ri in range(len(dd)):
            for ci, c in enumerate(cols):
                cell = tab.cell(ri + 1, ci); cell.text = ""
                r = cell.paragraphs[0].add_run(str(dd.iloc[ri][c])); r.font.size = Pt(7.5)
        if len(d) > a.report_max_rows or len(d.columns) > a.report_max_cols:
            para("Truncated for the page. Full data in {}.".format(n), italic=True)
    out = os.path.join(src, "ALL_RESULTS.docx")
    doc.save(out)
    say("written: {}".format(out))
    say("\n{} tables assembled. Re-run `report` after any analysis to refresh.".format(len(blocks)))


def cmd_scripts(a):
    tr, ev, th, eh = detect_scripts(a.project_root)
    say("Scanned {}\n".format(os.path.abspath(a.project_root)))
    say("TRAINER   : {}".format(tr or "NOT FOUND"))
    for f in th[1:]:
        say("    also: {}".format(f))
    say("EVALUATOR : {}".format(ev or "NOT FOUND"))
    for f in eh[1:]:
        say("    also: {}".format(f))
    if ev is None:
        say("\nNo file looks like a per-checkpoint evaluator. Pass it explicitly:")
        say("  python msc_analysis.py dro --eval-script YOUR_EVAL.py")
    else:
        f = eval_flags(ev)
        say("\nDetected flags for {}:".format(os.path.basename(ev)))
        say("  checkpoint -> --{}".format(f["ckpt"]))
        say("  splits     -> --{}".format(f["splits"]))
        say("  output     -> --{}".format(f["out"]))
        say("\nGenerated commands will use these names, so they will run as written.")
    return tr, ev


def n_classes_of(splits_dir, which="train.csv"):
    """Number of distinct labels in a split CSV, or None."""
    for cand in (splits_dir, os.path.join(splits_dir, "validated")):
        p = os.path.join(cand, which)
        if os.path.isfile(p):
            try:
                d = pd.read_csv(p)
                c = pick_col(d, LABEL_PREFS)
                if c is not None:
                    return int(d[c].nunique())
            except Exception:
                return None
    return None


def check_label_spaces(train_dir, cross_dir):
    """
    A model trained on K classes evaluated against a split with a different
    number of classes produces silent nonsense: the label indices simply do not
    mean the same thing. This caught a three-class model being scored against a
    two-class CN/AD split, where the model's MCI class was compared against the
    split's AD label and ROC-AUC came back NaN.
    """
    a = n_classes_of(train_dir, "train.csv")
    b = n_classes_of(cross_dir, "test.csv")
    if a is None or b is None:
        return True
    if a != b:
        say("")
        say("!" * 70)
        say("LABEL-SPACE MISMATCH. Training split has {} classes; the cross-cohort"
            .format(a))
        say("split has {}. Evaluating across these is meaningless: the label".format(b))
        say("indices do not refer to the same diagnoses.")
        say("  train: {}".format(train_dir))
        say("  cross: {}".format(cross_dir))
        say("Pass a --cross-splits-dir with {} classes, or omit it to skip the"
            .format(a))
        say("cross-cohort evaluation entirely.")
        say("!" * 70)
        say("")
        return False
    return True


def _bat(path, cmds):
    os.makedirs(os.path.dirname(path) or ".", exist_ok=True)
    with open(path, "w") as f:
        f.write("@echo off\r\n")
        for c in cmds:
            f.write(c + "\r\n")
    say("\nCommands written to {}\n".format(path))
    for c in cmds:
        say(c)


def cmd_dro(a):
    """Group DRO inventory, commands, and results table once predictions exist."""
    tr, ev, _, _ = detect_scripts(a.project_root)
    if a.eval_script == "AUTO":
        a.eval_script = ev or "evaluate.py"
    if a.train_script == "AUTO":
        a.train_script = tr or "train.py"
    say("Evaluator: {}\nTrainer:   {}\n".format(a.eval_script, a.train_script))
    if a.cross_splits_dir and not check_label_spaces(a.splits_dir,
                                                     a.cross_splits_dir):
        a.cross_splits_dir = None

    runs = []
    if os.path.isdir(a.dro_root):
        for d in sorted(os.listdir(a.dro_root)):
            full = os.path.join(a.dro_root, d)
            if not os.path.isdir(full):
                continue
            js = {}
            ts = os.path.join(full, "train_summary.json")
            if os.path.isfile(ts):
                try:
                    js = json.load(open(ts))
                except Exception:
                    pass
            m = re.search(r"fold[_-]?(\d+)", d.lower())
            runs.append(dict(run=d, dir=full,
                             has_ckpt=os.path.isfile(os.path.join(full, "best.pt")),
                             preds=glob.glob(os.path.join(full, "**", "predictions_test.csv"),
                                             recursive=True),
                             fold=int(m.group(1)) if m else None,
                             splits_dir=(js.get("args") or {}).get("splits_dir"),
                             val_best=js.get("best_score"), monitor=js.get("monitor")))
    if not runs:
        die("No DRO runs under {}".format(a.dro_root))

    say("Group DRO inventory\n" + "-" * 60)
    for r in runs:
        say("{:<34s} ckpt={} preds={} fold={}".format(
            r["run"], "Y" if r["has_ckpt"] else "n", len(r["preds"]), r["fold"]))
        if r["val_best"] is not None:
            say("    best_score={:.4f} ({})  <-- VALIDATION, DO NOT REPORT".format(
                r["val_best"], r["monitor"]))

    cmds = []
    for r in runs:
        if r["has_ckpt"] and not r["preds"]:
            sd = r["splits_dir"] or a.splits_dir
            cmds.append(eval_cmd(a.eval_script, r["dir"] + "\\best.pt", sd,
                                 r["dir"] + "\\eval_within"))
            if a.cross_splits_dir:
                cmds.append(eval_cmd(a.eval_script, r["dir"] + "\\best.pt",
                                     a.cross_splits_dir, r["dir"] + "\\eval_cross"))
    done = sorted({r["fold"] for r in runs if r["has_ckpt"] and r["fold"] is not None})
    base_sd = runs[0]["splits_dir"] or a.splits_dir or ""
    for i in [f for f in range(5) if f not in done]:
        sd = re.sub(r"fold_\d+", "fold_{}".format(i), base_sd)
        rd = os.path.join(a.dro_root, "dro_adni_r3d_18_fold_{}".format(i))
        cmds.append(
            'python "{ts}" --splits_dir "{sd}" --results_dir "{rd}" --backbone r3d_18 '
            '--pretrained --dropout 0.2 --group_dro --dro_eta 0.2 --dro_ema 0.9 '
            '--dro_q_floor 0.01 --epochs 60 --batch_size 2 --accum_steps 4 --lr 3e-4 '
            '--weight_decay 1e-4 --patch_size 96 --freeze_epochs 10 --early_patience 12 '
            '--monitor macro_f1 --label_smoothing 0.05 --focal_gamma 1.5 --grad_clip 1.0 '
            '--seed 1337'.format(ts=a.train_script, sd=sd, rd=rd))
        cmds.append(eval_cmd(a.eval_script, rd + "\\best.pt", sd, rd + "\\eval_within"))
        if a.cross_splits_dir:
            cmds.append(eval_cmd(a.eval_script, rd + "\\best.pt",
                                 a.cross_splits_dir, rd + "\\eval_cross"))
    if cmds:
        _bat(os.path.join(a.out, "run_dro.bat"), cmds)

    have = [r for r in runs if r["preds"]]
    if not have:
        say("\nNo DRO test predictions exist, so there is NO Group DRO result. Section 4.8")
        say("must stay pending. The validation figure is not a substitute.")
        return None
    rows = []
    edges = parse_edges(a.age_edges)
    for r in have:
        for p in r["preds"]:
            d = load_predictions(p)
            k = int(max(d.y.max(), d.pred.max())) + 1
            d = d.copy(); d["bin"] = assign_bins(d.age.values, edges)
            per = {b: macro_f1(s.y.values, s.pred.values, k)
                   for b in [e[2] for e in edges]
                   for s in [d[d.bin == b]] if len(s) >= a.min_bin}
            rows.append(dict(run=r["run"], fold=r["fold"],
                             macro_f1=macro_f1(d.y.values, d.pred.values, k),
                             balanced_acc=bal_acc(d.y.values, d.pred.values, k),
                             degenerate=degenerate(d.pred.values),
                             worst_bin=min(per.values()) if per else np.nan,
                             best_bin=max(per.values()) if per else np.nan,
                             gap=(max(per.values()) - min(per.values())) if len(per) > 1 else np.nan))
    df = pd.DataFrame(rows)
    write(df, a.out, "dro_results.csv", a.force)
    say("\nGroup DRO held-out test results:")
    show(df)
    say("\nIMPORTANT: check the splits_dir each evaluation actually used. If the")
    say("eval_within calls failed, these are CROSS-cohort numbers and the comparator")
    say("is the ADNI-to-OASIS baseline (macro-F1 0.378), NOT the within-cohort 0.680.")
    say("run_index.csv now carries a variant column so DRO rows never average into")
    say("the baseline ones.")
    return df


def cmd_splits(a):
    """
    Build age-bin-balanced OASIS three-class splits from the existing ones.

    Within each age bin independently, classes are undersampled to the smallest
    class in that bin. This is the same procedure the two-class arm already
    uses. Validation and test partitions are copied UNCHANGED, because
    evaluation must reflect the natural distribution; only training data is
    rebalanced. Fold membership is inherited, so the new arm is paired with the
    old one subject for subject and the comparison stays controlled.
    """
    src_root = a.source_splits
    if not src_root:
        die("Pass --source-splits pointing at the folder that CONTAINS fold_0..fold_4\n"
            "e.g. --source-splits \"C:\\\\Users\\\\todor\\\\Desktop\\\\Msc-AD\\\\data\\\\oasis\\\\splits\\\\res_128\"")
    out_root = a.balanced_out
    os.makedirs(out_root, exist_ok=True)
    edges = parse_edges(a.age_edges)
    rng = np.random.default_rng(SEED)

    fold_dirs = []
    for dp, dn, _fn in os.walk(src_root):
        for d in dn:
            if re.fullmatch(r"fold[_-]?\d+", d.lower()):
                fold_dirs.append(os.path.join(dp, d))
    fold_dirs = sorted(set(fold_dirs))
    if not fold_dirs:
        die("No fold_* directories under {}".format(src_root))
    say("Found {} fold directories under {}\n".format(len(fold_dirs), src_root))

    summary = []
    for fd in fold_dirs:
        # the split CSVs may sit in the fold dir or in a 'validated' subdir
        base = fd
        if not os.path.isfile(os.path.join(base, "train.csv")):
            for sub in ("validated", "splits", ""):
                cand = os.path.join(fd, sub)
                if os.path.isfile(os.path.join(cand, "train.csv")):
                    base = cand; break
        trp = os.path.join(base, "train.csv")
        if not os.path.isfile(trp):
            warn("no train.csv under {}, skipping".format(fd)); continue
        tr = pd.read_csv(trp)

        acol = pick_col(tr, AGE_PREFS)
        lcol = pick_col(tr, LABEL_PREFS)
        scol = pick_col(tr, SUBJ_PREFS)
        if acol is None or lcol is None:
            die("train.csv needs age and label columns. Found: {}".format(list(tr.columns)))

        tr = tr.copy()
        tr["_bin"] = assign_bins(tr[acol].values, edges)

        if a.balance_mode == "reweight":
            # Equalise the age-conditional prior WITHOUT discarding subjects:
            # each sample is weighted by the inverse of its class frequency
            # within its own age bin. This is the intervention Section 5.4 of
            # the dissertation recommends, and it keeps all 277 subjects.
            w = np.ones(len(tr), float)
            for b in tr["_bin"].unique():
                m = (tr["_bin"] == b).values
                cnt = tr.loc[m, lcol].value_counts()
                for c, k in cnt.items():
                    w[m & (tr[lcol] == c).values] = len(cnt) and (m.sum() / (len(cnt) * k))
            bal = tr.drop(columns=["_bin"]).copy()
            bal["sample_weight"] = w
            rel = os.path.relpath(base, src_root)
            dest = os.path.join(out_root, rel); os.makedirs(dest, exist_ok=True)
            bal.to_csv(os.path.join(dest, "train.csv"), index=False)
            for nm in ("val.csv", "valid.csv", "validation.csv", "test.csv"):
                sp = os.path.join(base, nm)
                if os.path.isfile(sp):
                    pd.read_csv(sp).to_csv(os.path.join(dest, nm), index=False)
            say("{}: {} subjects kept, sample_weight range {:.2f} to {:.2f}".format(
                rel, len(bal), w.min(), w.max()))
            summary.append(dict(fold=rel, before=len(tr), after=len(bal), dropped=0))
            continue

        keep = []
        for b in [e[2] for e in edges]:
            sub = tr[tr["_bin"] == b]
            if len(sub) == 0:
                continue
            counts = sub[lcol].value_counts()
            m = int(counts.min())
            if m == 0:
                continue
            for c in sorted(sub[lcol].unique()):
                keep.append(sub[sub[lcol] == c].sample(n=m, random_state=SEED))
        if not keep:
            warn("nothing to keep for {}".format(fd)); continue
        bal = pd.concat(keep, ignore_index=True).drop(columns=["_bin"])
        if scol:
            assert bal[scol].duplicated().sum() == 0, "duplicate subjects after balancing"

        rel = os.path.relpath(base, src_root)
        dest = os.path.join(out_root, rel)
        os.makedirs(dest, exist_ok=True)
        bal.to_csv(os.path.join(dest, "train.csv"), index=False)
        # validation and test copied unchanged: evaluation keeps the natural prior
        for nm in ("val.csv", "valid.csv", "validation.csv", "test.csv"):
            sp = os.path.join(base, nm)
            if os.path.isfile(sp):
                pd.read_csv(sp).to_csv(os.path.join(dest, nm), index=False)

        before = tr.groupby(["_bin", lcol]).size().unstack(fill_value=0)
        after = bal.assign(_bin=assign_bins(bal[acol].values, edges)) \
                   .groupby(["_bin", lcol]).size().unstack(fill_value=0)
        say("{}".format(rel))
        say("   before: {} subjects   per bin/class:\n{}".format(
            len(tr), before.to_string().replace("\n", "\n      ")))
        say("   after : {} subjects   per bin/class:\n{}".format(
            len(bal), after.to_string().replace("\n", "\n      ")))
        summary.append(dict(fold=rel, before=len(tr), after=len(bal),
                            dropped=len(tr) - len(bal)))

    df = pd.DataFrame(summary)
    write(df, a.out, "balanced_split_summary.csv", a.force)
    say("\nBalanced splits written under {}".format(os.path.abspath(out_root)))
    say("\nWhat changed: the class prior is now uniform WITHIN each age bin, not")
    say("merely uniform overall. Global class weights (already active in every")
    say("run) equalise the marginal prior; only this breaks the association")
    say("between age and diagnosis, which is the confound the study is about.")
    say("\nCost: {} of {} training subjects dropped across folds. That is the".format(
        int(df.dropped.sum()) if len(df) else 0,
        int(df.before.sum()) if len(df) else 0))
    say("price of the controlled comparison, and it is why the balanced arm's")
    say("absolute performance is expected to be LOWER. The point is the age")
    say("probe, not the accuracy.")
    first = ""
    for dp, dn, fn in os.walk(out_root):
        if "train.csv" in fn:
            first = os.path.abspath(dp); break
    say("\nNext, copy this line exactly:")
    say('  python msc_analysis.py balanced --balanced-splits-dir "{}"'.format(first))
    return df


def cmd_balanced(a):
    """
    Generate the commands for the missing age-bin-balanced OASIS three-class arm.

    Why this matters: at present every age-unbalanced cell is OASIS three-class,
    so balancing, cohort and arm are perfectly confounded. The permutation result
    shows those three cells rank top on age decodability; it cannot show that
    balancing is the reason. Training the balanced counterpart on ONE backbone,
    with everything else held constant, converts a correlation across cells into
    a controlled ablation.
    """
    tr, ev, _, _ = detect_scripts(a.project_root)
    train = a.train_script if a.train_script != "AUTO" else (tr or "train.py")
    evs = a.eval_script if a.eval_script != "AUTO" else (ev or "evaluate.py")
    say("Trainer:   {}\nEvaluator: {}\n".format(train, evs))
    if a.cross_splits_dir and not check_label_spaces(a.balanced_splits_dir,
                                                     a.cross_splits_dir):
        say("Cross-cohort evaluation lines will be OMITTED from the batch file.")
        say("The within-cohort runs and the age probe are unaffected, and they")
        say("are what the balanced arm exists to provide.\n")
        a.cross_splits_dir = None

    if not a.balanced_splits_dir:
        say("You need age-bin-balanced OASIS three-class splits first.")
        say("If your split builder has a balancing flag, regenerate with it enabled;")
        say("otherwise undersample each class to the smallest class WITHIN each age")
        say("bin independently, exactly as the two-class arm already does, keeping")
        say("fold membership unchanged so the comparison stays paired.\n")
        say("Then re-run:  python msc_analysis.py balanced --balanced-splits-dir <DIR>")
        return

    # Discover sibling fold directories, WITHOUT walking up the tree. Walking
    # up previously escaped into the data root and matched 50 directories,
    # including every dated backup and the ADNI splits.
    given = a.balanced_splits_dir
    if not os.path.isdir(given):
        die("--balanced-splits-dir does not exist: {}\n"
            "Use the path `splits` printed under 'Balanced splits written under'."
            .format(given))

    def _fold_base(d):
        """Return (fold_dir, csv_dir) for a directory that holds or contains train.csv."""
        if os.path.isfile(os.path.join(d, "train.csv")):
            return d
        for sub in ("validated", "splits"):
            if os.path.isfile(os.path.join(d, sub, "train.csv")):
                return os.path.join(d, sub)
        return None

    # If the given path is itself a fold, its siblings live in its parent.
    if re.fullmatch(r"fold[_-]?\d+", os.path.basename(given.rstrip("\\/")).lower()):
        parent = os.path.dirname(given.rstrip("\\/"))
    elif _fold_base(given):
        parent = os.path.dirname(os.path.dirname(given.rstrip("\\/")))
    else:
        parent = given

    fold_dirs = []
    if os.path.isdir(parent):
        for d in sorted(os.listdir(parent)):
            if re.fullmatch(r"fold[_-]?\d+", d.lower()):
                base = _fold_base(os.path.join(parent, d))
                if base:
                    fold_dirs.append(base)
    # one level deeper, e.g. <out>/5fold/fold_1
    if not fold_dirs and os.path.isdir(parent):
        for mid in sorted(os.listdir(parent)):
            mp = os.path.join(parent, mid)
            if not os.path.isdir(mp):
                continue
            for d in sorted(os.listdir(mp)):
                if re.fullmatch(r"fold[_-]?\d+", d.lower()):
                    base = _fold_base(os.path.join(mp, d))
                    if base:
                        fold_dirs.append(base)
    fold_dirs = [f for f in sorted(set(fold_dirs)) if "backup" not in f.lower()]
    if not fold_dirs:
        die("No fold directories with a train.csv beside {}.".format(given))
    if len(fold_dirs) > 8:
        die("Found {} fold directories under {}, which is more than one arm.\n"
            "Point --balanced-splits-dir at a single fold inside the output of "
            "`splits`.".format(len(fold_dirs), parent))
    say("Using {} fold directories:".format(len(fold_dirs)))
    for f in fold_dirs:
        say("  " + f)
    say("")

    cmds = []
    for sd in fold_dirs:
        m = re.search(r"fold[_-]?(\d+)", sd.lower())
        f = int(m.group(1)) if m else len(cmds)
        rd = os.path.join(a.results_dir, "oasis3_balanced_{}_fold_{}".format(a.backbone, f))
        # If the splits carry a sample_weight column they came from
        # --balance-mode reweight, and the trainer must be told to honour it.
        extra = ""
        try:
            if "sample_weight" in pd.read_csv(os.path.join(sd, "train.csv"),
                                              nrows=1).columns:
                extra = " --use_sample_weight"
        except Exception:
            pass
        cmds.append('python "{ts}" --splits_dir "{sd}" --results_dir "{rd}" '
                    '--backbone {bk} --pretrained --dropout 0.2 --epochs 60 '
                    '--batch_size 2 --accum_steps 4 --lr 1e-4 --weight_decay 1e-4 '
                    '--patch_size 96 --freeze_epochs 10 --early_patience 12 '
                    '--monitor macro_f1 --label_smoothing 0.05 --focal_gamma 1.5 '
                    '--grad_clip 1.0 --seed 1337{ex}'.format(ts=train, sd=sd, rd=rd,
                                                             bk=a.backbone, ex=extra))
        cmds.append(eval_cmd(evs, rd + "\\best.pt", sd, rd + "\\eval_within"))
        if a.cross_splits_dir:
            cmds.append(eval_cmd(evs, rd + "\\best.pt", a.cross_splits_dir,
                                 rd + "\\eval_cross"))
    _bat(os.path.join(a.out, "run_balanced_oasis3.bat"), cmds)
    say("\nLearning rate is 1e-4 to match the existing OASIS runs, so the ONLY thing")
    say("that differs from the unbalanced arm is the age-conditional balancing.")
    say("\nAfter these finish, re-run the age probe on both conditions. If age R2 drops")
    say("from roughly 0.5 to roughly 0 on the same architecture, same cohort and same")
    say("subjects, the confound is gone and the claim becomes causal.")


def cmd_ageshift(a):
    """Younger-skewed training splits against a fixed older evaluation set."""
    src = a.splits_dir
    if not os.path.isdir(src):
        die("--splits-dir not found: {}".format(src))
    def first(names):
        for n in names:
            p = os.path.join(src, n)
            if os.path.isfile(p):
                return p
    trp = first(["train.csv", "train_split.csv"]); tep = first(["test.csv", "test_split.csv"])
    vap = first(["val.csv", "valid.csv", "validation.csv"])
    if not trp or not tep:
        die("No train/test CSV in {}. Files: {}".format(src, os.listdir(src)))
    pool = pd.concat([pd.read_csv(trp)] + ([pd.read_csv(vap)] if vap else []) +
                     [pd.read_csv(tep)], ignore_index=True).drop_duplicates()
    acol = pick_col(pool, AGE_PREFS)
    lcol = pick_col(pool, LABEL_PREFS)
    scol = pick_col(pool, SUBJ_PREFS)
    if not all([acol, lcol, scol]):
        die("Split CSV needs age, label and subject-id columns. Found: {}\n"
            "Without a subject id a leakage-free split cannot be guaranteed."
            .format(list(pool.columns)))
    pool = pool.drop_duplicates(subset=[scol]).reset_index(drop=True)
    cut = float(np.quantile(pool[acol].values, 1 - a.eval_frac))
    older = pool[pool[acol] >= cut]
    npc = int(older[lcol].value_counts().min())
    ev = pd.concat([older[older[lcol] == c].sample(n=npc, random_state=SEED)
                    for c in sorted(older[lcol].unique())], ignore_index=True)
    ev_ids = set(ev[scol].astype(str))
    rest = pool[~pool[scol].astype(str).isin(ev_ids)].reset_index(drop=True)
    assert not (set(rest[scol].astype(str)) & ev_ids)
    say("Pool {} subjects. Fixed eval set: {} aged >= {:.1f}, {} per class. "
        "Training pool: {}.".format(len(pool), len(ev), cut, npc, len(rest)))
    outdir = os.path.join(a.out, "ageshift_splits"); os.makedirs(outdir, exist_ok=True)
    med = float(np.median(rest[acol].values))
    young, old = rest[rest[acol] < med], rest[rest[acol] >= med]
    cmds = []
    for sk in [float(x) for x in a.skews.split(",")]:
        tot = min(len(young), len(old)) * 2
        ny = min(int(round(sk * tot)), len(young)); no = min(tot - ny, len(old))
        part = pd.concat([young.sample(ny, random_state=SEED),
                          old.sample(no, random_state=SEED)], ignore_index=True)
        m = int(part[lcol].value_counts().min())
        part = pd.concat([part[part[lcol] == c].sample(m, random_state=SEED)
                          for c in sorted(part[lcol].unique())],
                         ignore_index=True).reset_index(drop=True)
        hold = part.sample(frac=0.1, random_state=SEED); core = part.drop(hold.index)
        assert not (set(core[scol].astype(str)) & set(hold[scol].astype(str)))
        assert not (set(core[scol].astype(str)) & ev_ids)
        cd = os.path.join(outdir, "skew_{:.0f}".format(sk * 100)); os.makedirs(cd, exist_ok=True)
        core.to_csv(os.path.join(cd, "train.csv"), index=False)
        hold.to_csv(os.path.join(cd, "val.csv"), index=False)
        ev.to_csv(os.path.join(cd, "test.csv"), index=False)
        say("  skew {:>4.0f}% young: train {:3d}, val {:2d}, mean train age {:.1f} "
            "(eval {:.1f}, shift {:+.1f} yr)".format(sk * 100, len(core), len(hold),
                                                     core[acol].mean(), ev[acol].mean(),
                                                     core[acol].mean() - ev[acol].mean()))
        rd = os.path.join(a.results_dir, "ageshift_{}_skew{:.0f}".format(a.backbone, sk * 100))
        cmds.append('python "{}" --splits_dir "{}" --results_dir "{}" --backbone {} '
                    '--pretrained --epochs 60 --batch_size 2 --accum_steps 4 --lr 3e-4 '
                    '--patch_size 96 --freeze_epochs 10 --early_patience 12 '
                    '--monitor macro_f1 --seed 1337'.format(a.train_script, cd, rd, a.backbone))
        cmds.append(eval_cmd(a.eval_script, rd + "\\best.pt", cd, rd + "\\eval_fixedold"))
    _bat(os.path.join(a.out, "run_ageshift.bat"), cmds)
    say("\nThe evaluation set is IDENTICAL across conditions, so any change is "
        "attributable to the training age profile alone.")


def cmd_seeds(a):
    if a.collect:
        files = [f for f in find_predictions(a.roots) if "seed" in f.lower()]
        if not files:
            die("No seed runs found.")
        rows = []
        for f in files:
            m = re.search(r"seed[_-]?(\d+)", f.lower())
            d = load_predictions(f); k = int(max(d.y.max(), d.pred.max())) + 1
            rows.append(dict(seed=int(m.group(1)) if m else -1,
                             macro_f1=macro_f1(d.y.values, d.pred.values, k),
                             balanced_acc=bal_acc(d.y.values, d.pred.values, k), path=f))
        df = pd.DataFrame(rows)
        write(df, a.out, "seed_sensitivity.csv", a.force)
        show(df, ["seed", "macro_f1", "balanced_acc"])
        say("\nmacro-F1 sd across seeds: {:.3f}".format(df.macro_f1.std(ddof=1)))
        say("Compare against the across-FOLD sd in cv_performance.csv. If they are")
        say("comparable, fold variance is a fair stand-in and can be declared as such.")
        return df
    cmds = []
    for s in [int(x) for x in a.seed_list.split(",")]:
        rd = os.path.join(a.results_dir, "seed_{}_{}_fold{}".format(s, a.backbone, a.fold))
        cmds.append('python "{}" --splits_dir "{}" --results_dir "{}" --backbone {} '
                    '--pretrained --epochs 60 --batch_size 2 --accum_steps 4 --lr 3e-4 '
                    '--patch_size 96 --freeze_epochs 10 --early_patience 12 '
                    '--monitor macro_f1 --seed {}'.format(
                        a.train_script, a.splits_dir, rd, a.backbone, s))
        cmds.append(eval_cmd(a.eval_script, rd + "\\best.pt", a.splits_dir,
                             rd + "\\eval_within"))
    _bat(os.path.join(a.out, "run_seeds.bat"), cmds)
    say("\nThen: python msc_analysis.py seeds --collect")


# =========================================================================
# CLI
# =========================================================================

CPU_CMDS = ["index", "cv", "paired", "perbin", "bins", "confage", "thresh",
            "agegap", "degen", "domain", "calib", "figures"]


def cmd_all(a):
    done, failed = [], []
    for name in CPU_CMDS:
        banner(name.upper())
        try:
            globals()["cmd_" + name](a); done.append(name)
        except SystemExit:
            raise
        except Exception as e:
            warn("{} failed: {}".format(name, e)); failed.append((name, str(e)))
    banner("SUMMARY")
    say("Completed: {}".format(", ".join(done) or "none"))
    if failed:
        say("Failed:")
        for n, e in failed:
            say("  {}: {}".format(n, e))
    say("\nOutputs in {}".format(os.path.abspath(a.out)))
    say("\nStill needing a GPU or the volumes: config, embed, reorient, roi.")
    say("Still needing training runs: dro (folds 1-4), ageshift, seeds.")


def main():
    ap = argparse.ArgumentParser(description=__doc__,
                                 formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("command", nargs="?", default="all",
                    choices=CPU_CMDS + ["all", "config", "embed", "probe", "reorient",
                                        "roi", "cam", "scripts", "dro", "ageshift",
                                        "seeds", "balanced", "splits", "report"])
    ap.add_argument("--roots", nargs="+", default=None)
    ap.add_argument("--out", default=DEFAULT_OUT)
    ap.add_argument("--force", action="store_true")
    ap.add_argument("--age-edges", default=None, help="e.g. 70,80")
    ap.add_argument("--min-bin", type=int, default=10)
    ap.add_argument("--reps", type=int, default=N_BOOT)
    ap.add_argument("--perm", type=int, default=2000)
    ap.add_argument("--oas1-pattern", default=r"OAS1", help="regex marking OASIS-1 ids")
    ap.add_argument("--project-root", default=".")
    ap.add_argument("--model-module", default=None)
    ap.add_argument("--train-script", default="AUTO")
    ap.add_argument("--eval-script", default="AUTO")
    ap.add_argument("--results-dir", default="results_extra")
    ap.add_argument("--backbone", default="r3d_18")
    ap.add_argument("--fold", type=int, default=0)
    ap.add_argument("--splits-dir",
                    default=r"C:\Users\todor\Desktop\Msc-AD\data\splits_adni_cnad\fold_0\validated")
    ap.add_argument("--cross-splits-dir",
                    default=r"C:\Users\todor\Desktop\Msc-AD\data\oasis\splits\res_128_cnad\all\validated")
    ap.add_argument("--report-max-rows", type=int, default=60)
    ap.add_argument("--report-max-cols", type=int, default=12)
    ap.add_argument("--dro-root", default="results_dro")
    ap.add_argument("--balance-mode", default="undersample",
                    choices=["undersample", "reweight"],
                    help="undersample matches the two-class protocol exactly; "
                         "reweight keeps every subject and writes a "
                         "sample_weight column instead")
    ap.add_argument("--source-splits", default=None,
                    help="folder CONTAINING fold_0..fold_4 of the unbalanced arm")
    ap.add_argument("--balanced-out", default="splits_oasis3_balanced",
                    help="where to write the balanced splits")
    ap.add_argument("--balanced-splits-dir", default=None,
                    help="age-bin-balanced OASIS 3-class splits, fold_0 path")
    ap.add_argument("--reorient-csv", default="reorient_full.csv")
    ap.add_argument("--eval-frac", type=float, default=0.30)
    ap.add_argument("--skews", default="0.2,0.4,0.6,0.8")
    ap.add_argument("--seed-list", default="1337,7,42,2024,31337")
    ap.add_argument("--collect", action="store_true")
    ap.add_argument("--probe-sizes", default="277,132,64")
    ap.add_argument("--embed-filter", default=None,
                    help="regex; only probe matching run directories")
    ap.add_argument("--template-dir",
                    default=r"C:\Users\todor\Desktop\Msc-AD\data\oasis\processed\res_128")
    ap.add_argument("--n-template", type=int, default=100)
    ap.add_argument("--roi", default="52,78,58,84,44,84")
    # attribution
    ap.add_argument("--checkpoint", default=None, help="best.pt to attribute from")
    ap.add_argument("--cam-split", default=None, help="split CSV listing scan paths")
    ap.add_argument("--cam-methods", default="hirescam,gradcam,gradcampp")
    ap.add_argument("--cam-layer", default=None, help="e.g. layer4; default auto")
    ap.add_argument("--cam-target", default="true", choices=["true", "pred"])
    ap.add_argument("--cam-slices", type=int, default=7)
    ap.add_argument("--cam-alpha", type=float, default=0.45)
    ap.add_argument("--cam-thresh", type=float, default=0.35,
                    help="hide attribution below this, 0-1; raise for cleaner figures")
    ap.add_argument("--cam-cmap", default="inferno")
    ap.add_argument("--cam-dpi", type=int, default=220)
    ap.add_argument("--cam-n", type=int, default=120, help="subjects to average")
    ap.add_argument("--cam-examples", type=int, default=4)
    ap.add_argument("--cam-sanity", action="store_true", default=True)
    ap.add_argument("--cam-sanity-n", type=int, default=20)
    ap.add_argument("--oas1-only", action="store_true", default=True,
                    help="group maps need a shared space; only OASIS-1 has one")
    ap.add_argument("--patch-size", type=int, default=96)
    ap.add_argument("--cpu", action="store_true")
    a = ap.parse_args()

    os.makedirs(a.out, exist_ok=True)
    if a.roots is None:
        found = [r for r in DEFAULT_ROOTS if os.path.isdir(r)]
        a.roots = found or ["."]
        n = len(find_predictions(a.roots))
        if n == 0 and a.roots != ["."]:
            a.roots = ["."]; n = len(find_predictions(a.roots))
        say("Roots: {}   ({} prediction files)\n".format(", ".join(a.roots), n))

    if a.command == "all":
        cmd_all(a)
    else:
        globals()["cmd_" + a.command](a)


if __name__ == "__main__":
    main()
